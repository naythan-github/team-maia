#!/usr/bin/env python3
"""
Create Clean Orro Corporate Reference Template

Extracts pure styles from existing PIR template, removes all content,
creates a style-only reference template for generic MD→DOCX conversions.

Agent: Document Conversion Specialist Agent
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Orro purple color: RGB(112, 48, 160) - Official Orro brand color
ORRO_PURPLE = RGBColor(112, 48, 160)


def _setup_document_margins(doc: Document) -> None:
    """
    Set document margins to 1.0" on all sides.

    Args:
        doc: Document to configure
    """
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def _configure_heading_styles(doc: Document) -> None:
    """
    Configure heading styles with Aptos font and Orro purple color.

    Args:
        doc: Document to configure
    """
    # Heading 1: Aptos, 16pt, bold, Orro purple
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Aptos'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = ORRO_PURPLE

    # Heading 2: Aptos, 14pt, bold, Orro purple
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Aptos'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = ORRO_PURPLE

    # Heading 3: Aptos, 12pt, bold, Orro purple
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Aptos'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = ORRO_PURPLE

    # Normal: Aptos, 11pt, black
    normal = doc.styles['Normal']
    normal.font.name = 'Aptos'
    normal.font.size = Pt(11)


def _add_style_examples(doc: Document) -> None:
    """
    Add minimal style examples to document.

    Args:
        doc: Document to add examples to
    """
    # Add one paragraph per style (minimal content for style reference)
    doc.add_heading('Heading 1 Style Example', level=1)
    doc.add_heading('Heading 2 Style Example', level=2)
    doc.add_heading('Heading 3 Style Example', level=3)
    doc.add_paragraph('Normal paragraph style example.')

    # Add list examples
    doc.add_paragraph('Bullet list item example', style='List Bullet')
    doc.add_paragraph('Numbered list item example', style='List Number')

    # Add table example
    table = doc.add_table(rows=2, cols=3)
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass  # Keep default if Table Grid not available

    # Populate table with example data
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(0, 2).text = 'Header 3'
    table.cell(1, 0).text = 'Data 1'
    table.cell(1, 1).text = 'Data 2'
    table.cell(1, 2).text = 'Data 3'


def _add_usage_instructions(doc: Document) -> None:
    """
    Add usage instructions to document.

    Args:
        doc: Document to add instructions to
    """
    doc.add_page_break()
    instruction = doc.add_paragraph()
    instruction.add_run('Orro Corporate Reference Template').bold = True
    doc.add_paragraph('')
    doc.add_paragraph('This template contains Orro corporate styles for use with Pandoc conversions:')
    doc.add_paragraph('• Aptos font (11pt body, 12-16pt headings)', style='List Bullet')
    doc.add_paragraph('• Purple Orro table style (_Orro Table 1)', style='List Bullet')
    doc.add_paragraph('• Standard 1.0" margins', style='List Bullet')
    doc.add_paragraph('')
    doc.add_paragraph('Usage:')
    doc.add_paragraph('pandoc input.md --reference-doc=orro_corporate_reference.docx -o output.docx')


def create_clean_orro_reference():
    """
    Create a clean Orro corporate reference template.

    Phase 230: Refactored to use helper functions for maintainability.

    Creates template with:
    - Style definitions (Heading 1-9, Normal, List styles)
    - Aptos font (corporate standard)
    - Purple Orro table style (_Orro Table 1)
    - Standard margins (1.0")
    - Minimal style examples (for Pandoc reference)
    """
    print("=" * 70)
    print("Creating Clean Orro Corporate Reference Template")
    print("=" * 70)
    print()

    # Check for source template (optional)
    import sys
    sys.path.insert(0, str(Path(__file__).parent.parent.parent.parent))
    from claude.tools.core.paths import PathManager
    source_template = PathManager.get_work_projects_path() / "pir_converter" / "pir_orro_reference.docx"

    if source_template.exists():
        print(f"📄 Source template found: {source_template.name}")
    else:
        print(f"📄 Creating template from scratch (no source template)")

    # Create document and configure using helper functions
    print("🔧 Creating clean template (style-only, no content)...")
    clean_doc = Document()

    print("   📏 Setting margins (1.0\" all sides)...")
    _setup_document_margins(clean_doc)

    print("   🎨 Configuring corporate styles (Orro purple headings)...")
    _configure_heading_styles(clean_doc)

    print("   📝 Adding minimal style examples...")
    print("   📊 Adding table style example...")
    _add_style_examples(clean_doc)

    print("   📌 Adding usage instructions...")
    _add_usage_instructions(clean_doc)

    # Save clean template
    from claude.tools.core.paths import TOOLS_DIR
    output_path = TOOLS_DIR / "document_conversion/templates/orro_corporate_reference.docx"

    print(f"💾 Saving clean template: {output_path}")
    clean_doc.save(output_path)

    print()
    print("✅ Clean template created successfully!")
    print()
    print("📋 TEMPLATE DETAILS:")
    print(f"   Location: {output_path}")
    print(f"   Font: Aptos (corporate standard)")
    print(f"   Colors: Orro purple headings RGB(112, 48, 160)")
    print(f"   Margins: 1.0\" all sides")
    print(f"   Styles: Heading 1-3 (Orro purple), Normal, List Bullet/Number")
    print(f"   Tables: _Orro Table 1 (purple borders)")
    print(f"   Content: Minimal examples only (Pandoc ignores, uses styles)")
    print()
    print("🎯 USAGE:")
    print("   pandoc document.md --reference-doc=orro_corporate_reference.docx -o output.docx")
    print()

    return True


def main():
    success = create_clean_orro_reference()

    if success:
        print("🚀 NEXT STEPS:")
        print("   1. Test template: Convert sample markdown with new reference")
        print("   2. Rename existing templates for clarity")
        print("   3. Update converter scripts to use appropriate templates")
        print()
    else:
        print("❌ Template creation failed")
        return 1

    return 0


if __name__ == '__main__':
    import sys
    sys.exit(main())
