"""
Tests for create_clean_orro_template.py helper functions.

TDD: Phase 4 refactoring - decompose create_clean_orro_reference (149 lines)
"""

import pytest
import sys
from pathlib import Path
from unittest.mock import MagicMock, patch

# Add parent to path
sys.path.insert(0, str(Path(__file__).parent.parent))


class TestHelperFunctionsExist:
    """RED: Verify helper functions exist after refactoring."""

    def test_setup_document_margins_exists(self):
        """Helper for setting document margins."""
        from create_clean_orro_template import _setup_document_margins
        assert callable(_setup_document_margins)

    def test_configure_heading_styles_exists(self):
        """Helper for configuring heading styles."""
        from create_clean_orro_template import _configure_heading_styles
        assert callable(_configure_heading_styles)

    def test_add_style_examples_exists(self):
        """Helper for adding style example content."""
        from create_clean_orro_template import _add_style_examples
        assert callable(_add_style_examples)

    def test_add_usage_instructions_exists(self):
        """Helper for adding usage instructions."""
        from create_clean_orro_template import _add_usage_instructions
        assert callable(_add_usage_instructions)


class TestSetupDocumentMargins:
    """Test _setup_document_margins helper."""

    def test_sets_all_margins_to_one_inch(self):
        """Should set all margins to 1.0 inch."""
        from create_clean_orro_template import _setup_document_margins
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        _setup_document_margins(doc)

        section = doc.sections[0]
        assert section.top_margin == Inches(1.0)
        assert section.bottom_margin == Inches(1.0)
        assert section.left_margin == Inches(1.0)
        assert section.right_margin == Inches(1.0)


class TestConfigureHeadingStyles:
    """Test _configure_heading_styles helper."""

    def test_sets_heading_fonts_to_aptos(self):
        """Should set heading fonts to Aptos."""
        from create_clean_orro_template import _configure_heading_styles
        from docx import Document

        doc = Document()
        _configure_heading_styles(doc)

        h1 = doc.styles['Heading 1']
        assert h1.font.name == 'Aptos'

    def test_sets_orro_purple_color(self):
        """Should set heading color to Orro purple."""
        from create_clean_orro_template import _configure_heading_styles
        from docx import Document
        from docx.shared import RGBColor

        doc = Document()
        _configure_heading_styles(doc)

        h1 = doc.styles['Heading 1']
        # Orro purple: RGB(112, 48, 160)
        assert h1.font.color.rgb == RGBColor(112, 48, 160)


class TestAddStyleExamples:
    """Test _add_style_examples helper."""

    def test_adds_heading_examples(self):
        """Should add heading examples to document."""
        from create_clean_orro_template import _add_style_examples
        from docx import Document

        doc = Document()
        _add_style_examples(doc)

        # Should have at least 3 headings + paragraphs
        assert len(doc.paragraphs) >= 6


class TestAddUsageInstructions:
    """Test _add_usage_instructions helper."""

    def test_adds_pandoc_usage_example(self):
        """Should include pandoc usage command."""
        from create_clean_orro_template import _add_usage_instructions
        from docx import Document

        doc = Document()
        _add_usage_instructions(doc)

        # Check that pandoc is mentioned
        all_text = ' '.join([p.text for p in doc.paragraphs])
        assert 'pandoc' in all_text.lower()
