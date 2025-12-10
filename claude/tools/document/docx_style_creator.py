#!/usr/bin/env python3
"""
Word Document Style Creator.

Detects formatting patterns in Word documents and creates named styles
from those patterns. Enables converting direct formatting to style-based
formatting for cleaner, more maintainable documents.

OOXML Structure Reference:
- Bold/Italic: w:rPr/w:b, w:rPr/w:i in run properties
- Lists: w:pPr/w:numPr referencing numbering.xml
- Fonts: w:rPr/w:rFonts, w:rPr/w:sz
- Style definitions: word/styles.xml w:style elements
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class FormatPattern:
    """Represents a detected formatting pattern."""
    pattern_type: str  # 'bold', 'italic', 'list', 'font', etc.
    count: int = 0
    examples: list = field(default_factory=list)
    properties: dict = field(default_factory=dict)


def detect_formatting_patterns(doc_path: Path | str) -> dict[str, dict[str, Any]]:
    """
    Detect formatting patterns in a Word document.

    Scans the document for direct formatting (not style-based) including:
    - Bold text runs
    - Italic text runs
    - List/bullet paragraphs
    - Custom fonts

    Args:
        doc_path: Path to the Word document

    Returns:
        Dictionary mapping pattern types to their details:
        {
            'bold': {'count': 10, 'examples': ['text1', 'text2']},
            'italic': {'count': 5, 'examples': [...]},
            'list': {'count': 20, 'examples': [...]},
        }
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required for style detection")

    doc = Document(str(doc_path))

    patterns = {
        'bold': {'count': 0, 'examples': []},
        'italic': {'count': 0, 'examples': []},
        'list': {'count': 0, 'examples': []},
    }

    # Scan paragraphs for list formatting
    for para in doc.paragraphs:
        is_list = False

        # Check for list/bullet formatting via numPr (direct formatting)
        if para._p.pPr is not None:
            num_pr = para._p.pPr.find(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr'
            )
            if num_pr is not None:
                is_list = True

        # Also check for list-based styles (style-based formatting)
        if para.style and para.style.name:
            style_name = para.style.name.lower()
            if 'list' in style_name or 'bullet' in style_name or 'number' in style_name:
                is_list = True

        if is_list:
            patterns['list']['count'] += 1
            if len(patterns['list']['examples']) < 3:
                text = para.text[:50] if para.text else ''
                patterns['list']['examples'].append(text)

        # Scan runs for character formatting
        for run in para.runs:
            # Check for bold (direct formatting, not from style)
            if run.bold is True:
                patterns['bold']['count'] += 1
                if len(patterns['bold']['examples']) < 3:
                    patterns['bold']['examples'].append(run.text[:30] if run.text else '')

            # Check for italic (direct formatting)
            if run.italic is True:
                patterns['italic']['count'] += 1
                if len(patterns['italic']['examples']) < 3:
                    patterns['italic']['examples'].append(run.text[:30] if run.text else '')

    # Also scan tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.bold is True:
                            patterns['bold']['count'] += 1
                        if run.italic is True:
                            patterns['italic']['count'] += 1

    return patterns


def create_style(
    doc: "Document",
    name: str,
    style_type: str = "paragraph",
    based_on: str | None = None,
    font_name: str | None = None,
    font_size: int | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: tuple | None = None,
) -> Any:
    """
    Create a new style in the document.

    Args:
        doc: python-docx Document object
        name: Name for the new style
        style_type: 'paragraph' or 'character'
        based_on: Name of style to base this on (for paragraph styles)
        font_name: Font family name
        font_size: Font size in points
        bold: Whether text should be bold
        italic: Whether text should be italic
        color: RGB color tuple (r, g, b)

    Returns:
        The created style object
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required for style creation")

    # Map style type string to enum
    if style_type == "paragraph":
        wd_type = WD_STYLE_TYPE.PARAGRAPH
    elif style_type == "character":
        wd_type = WD_STYLE_TYPE.CHARACTER
    else:
        raise ValueError(f"Unknown style type: {style_type}")

    # Check if style already exists
    try:
        existing = doc.styles[name]
        return existing
    except KeyError:
        pass

    # Create the style
    style = doc.styles.add_style(name, wd_type)

    # Set base style for paragraph styles
    if based_on and style_type == "paragraph":
        try:
            style.base_style = doc.styles[based_on]
        except KeyError:
            pass  # Base style doesn't exist

    # Apply font settings
    if font_name:
        style.font.name = font_name

    if font_size:
        style.font.size = Pt(font_size)

    if bold is not None:
        style.font.bold = bold

    if italic is not None:
        style.font.italic = italic

    if color:
        style.font.color.rgb = RGBColor(*color)

    return style


def apply_style_to_pattern(
    doc: "Document",
    pattern_type: str,
    style_name: str,
) -> int:
    """
    Apply a character style to runs matching a formatting pattern.

    Args:
        doc: python-docx Document object
        pattern_type: Type of pattern to match ('bold', 'italic')
        style_name: Name of style to apply

    Returns:
        Number of runs the style was applied to
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required")

    count = 0

    # Get the style
    try:
        style = doc.styles[style_name]
    except KeyError:
        raise ValueError(f"Style '{style_name}' not found in document")

    # Scan paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            matched = False

            if pattern_type == "bold" and run.bold is True:
                matched = True
            elif pattern_type == "italic" and run.italic is True:
                matched = True

            if matched:
                # Apply the character style
                run.style = style
                # Clear direct formatting since style now handles it
                if pattern_type == "bold":
                    run.bold = None  # Inherit from style
                elif pattern_type == "italic":
                    run.italic = None
                count += 1

    # Also process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        matched = False

                        if pattern_type == "bold" and run.bold is True:
                            matched = True
                        elif pattern_type == "italic" and run.italic is True:
                            matched = True

                        if matched:
                            run.style = style
                            if pattern_type == "bold":
                                run.bold = None
                            elif pattern_type == "italic":
                                run.italic = None
                            count += 1

    return count


def process_document(
    input_path: Path | str,
    output_path: Path | str,
    create_styles: list[str] | None = None,
) -> dict[str, Any]:
    """
    Full workflow: detect patterns, create styles, apply to content.

    Args:
        input_path: Path to input document
        output_path: Path for output document
        create_styles: List of style names to create. Supported:
            - 'Strong': Bold character style
            - 'Emphasis': Italic character style

    Returns:
        Statistics about the operation:
        {
            'styles_created': ['Strong'],
            'patterns_applied': {'bold': 5, 'italic': 3},
            'patterns_detected': {...}
        }
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required")

    input_path = Path(input_path)
    output_path = Path(output_path)

    # Detect patterns
    patterns = detect_formatting_patterns(input_path)

    # Open document for modification
    doc = Document(str(input_path))

    result = {
        'styles_created': [],
        'patterns_applied': {},
        'patterns_detected': patterns,
    }

    # Create and apply requested styles
    if create_styles:
        for style_name in create_styles:
            if style_name == 'Strong':
                # Create bold character style
                create_style(doc, 'Strong', style_type='character', bold=True)
                result['styles_created'].append('Strong')

                # Apply to bold runs
                count = apply_style_to_pattern(doc, 'bold', 'Strong')
                result['patterns_applied']['bold'] = count

            elif style_name == 'Emphasis':
                # Create italic character style
                create_style(doc, 'Emphasis', style_type='character', italic=True)
                result['styles_created'].append('Emphasis')

                # Apply to italic runs
                count = apply_style_to_pattern(doc, 'italic', 'Emphasis')
                result['patterns_applied']['italic'] = count

    # Save output
    doc.save(str(output_path))

    return result


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: docx_style_creator.py <document.docx> [output.docx]")
        print("\nDetects formatting patterns and creates styles.")
        sys.exit(1)

    input_path = Path(sys.argv[1])

    if not input_path.exists():
        print(f"Error: {input_path} not found")
        sys.exit(1)

    # Just detect patterns if no output specified
    if len(sys.argv) < 3:
        patterns = detect_formatting_patterns(input_path)
        print("\nDetected Formatting Patterns:")
        print("-" * 40)
        for name, data in patterns.items():
            if data['count'] > 0:
                print(f"\n{name.upper()}: {data['count']} occurrences")
                if data.get('examples'):
                    print("  Examples:")
                    for ex in data['examples'][:3]:
                        print(f"    - {ex}")
    else:
        output_path = Path(sys.argv[2])
        result = process_document(input_path, output_path, create_styles=['Strong', 'Emphasis'])
        print(f"\nProcessed: {input_path}")
        print(f"Output: {output_path}")
        print(f"Styles created: {result['styles_created']}")
        print(f"Patterns applied: {result['patterns_applied']}")
