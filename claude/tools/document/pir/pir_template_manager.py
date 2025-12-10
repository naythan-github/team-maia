#!/usr/bin/env python3
"""
Post-Incident Review (PIR) Template Manager
Security Specialist Agent - Template System

Manages PIR document templates for security incidents, enabling:
- Template extraction from completed PIRs
- Template application to new incidents
- Section customization and reordering
- Forensic analysis integration
"""

import os
import shutil
from pathlib import Path
from datetime import datetime
from docx import Document

class PIRTemplateManager:
    """Manages Post-Incident Review document templates"""

    def __init__(self, maia_root: str = None):
        """Initialize PIR Template Manager

        Args:
            maia_root: Path to Maia root directory (defaults to ~/git/maia)
        """
        if maia_root is None:
            maia_root = os.path.expanduser("~/git/maia")

        self.maia_root = Path(maia_root)
        self.template_dir = self.maia_root / "claude" / "tools" / "security" / "pir_templates"
        self.template_dir.mkdir(parents=True, exist_ok=True)

    def save_as_template(self, source_docx: str, template_name: str,
                        description: str = "", metadata: dict = None):
        """Save a completed PIR as a reusable template

        Args:
            source_docx: Path to source DOCX file (completed PIR)
            template_name: Name for the template (e.g., "pir_credential_stuffing_template")
            description: Description of when to use this template
            metadata: Additional metadata (incident_type, sections_included, etc.)

        Returns:
            Path to saved template
        """
        source_path = Path(source_docx)
        if not source_path.exists():
            raise FileNotFoundError(f"Source document not found: {source_docx}")

        # Template file paths
        template_file = self.template_dir / f"{template_name}.docx"
        metadata_file = self.template_dir / f"{template_name}.json"

        # Copy DOCX file
        shutil.copy2(source_path, template_file)

        # Create metadata
        template_metadata = {
            "template_name": template_name,
            "description": description,
            "created_date": datetime.now().isoformat(),
            "source_file": str(source_path),
            "incident_type": metadata.get("incident_type") if metadata else "generic",
            "sections": self._extract_sections(template_file),
            "custom_metadata": metadata or {}
        }

        # Save metadata
        import json
        with open(metadata_file, 'w') as f:
            json.dump(template_metadata, f, indent=2)

        print(f"✅ Template saved: {template_file}")
        print(f"✅ Metadata saved: {metadata_file}")

        return template_file

    def _extract_sections(self, docx_path: Path) -> list:
        """Extract section headings from DOCX file

        Args:
            docx_path: Path to DOCX file

        Returns:
            List of section headings
        """
        doc = Document(docx_path)
        sections = []

        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                sections.append({
                    'level': int(para.style.name.replace('Heading ', '')),
                    'text': para.text,
                    'style': para.style.name
                })

        return sections

    def list_templates(self) -> list:
        """List all available PIR templates

        Returns:
            List of template metadata dictionaries
        """
        import json
        templates = []

        for metadata_file in self.template_dir.glob("*.json"):
            with open(metadata_file, 'r') as f:
                templates.append(json.load(f))

        return templates

    def create_from_template(self, template_name: str, output_path: str,
                           incident_data: dict) -> Path:
        """Create a new PIR from template with incident-specific data

        Args:
            template_name: Name of template to use
            output_path: Where to save the new PIR
            incident_data: Dictionary of incident-specific values
                {
                    'ticket_number': '4184007',
                    'customer_name': 'NQLC',
                    'incident_date': '2025-11-19',
                    'incident_type': 'Account Compromise',
                    'severity': 'SEV1',
                    ...
                }

        Returns:
            Path to created PIR document
        """
        template_file = self.template_dir / f"{template_name}.docx"

        if not template_file.exists():
            raise FileNotFoundError(f"Template not found: {template_name}")

        # Copy template to output location
        output_path = Path(output_path)
        shutil.copy2(template_file, output_path)

        # Replace placeholders in the document
        doc = Document(output_path)

        # Define placeholder mappings
        placeholders = {
            '{{TICKET_NUMBER}}': incident_data.get('ticket_number', 'XXXXXX'),
            '{{CUSTOMER_NAME}}': incident_data.get('customer_name', 'Customer Name'),
            '{{INCIDENT_DATE}}': incident_data.get('incident_date', 'YYYY-MM-DD'),
            '{{INCIDENT_TYPE}}': incident_data.get('incident_type', 'Security Incident'),
            '{{SEVERITY}}': incident_data.get('severity', 'SEVX'),
            '{{ANALYST_NAME}}': incident_data.get('analyst_name', 'Security Team'),
            '{{REPORT_DATE}}': datetime.now().strftime('%Y-%m-%d'),
        }

        # Replace in paragraphs
        for para in doc.paragraphs:
            for placeholder, value in placeholders.items():
                if placeholder in para.text:
                    for run in para.runs:
                        run.text = run.text.replace(placeholder, value)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for placeholder, value in placeholders.items():
                            if placeholder in para.text:
                                for run in para.runs:
                                    run.text = run.text.replace(placeholder, value)

        doc.save(output_path)

        print(f"✅ PIR created from template: {output_path}")
        return output_path


def main():
    """CLI interface for PIR Template Manager"""
    import argparse

    parser = argparse.ArgumentParser(
        description="Post-Incident Review Template Manager"
    )

    subparsers = parser.add_subparsers(dest='command', help='Command to execute')

    # Save template command
    save_parser = subparsers.add_parser('save', help='Save a PIR as template')
    save_parser.add_argument('source', help='Source DOCX file path')
    save_parser.add_argument('name', help='Template name')
    save_parser.add_argument('--description', help='Template description')
    save_parser.add_argument('--incident-type', help='Incident type (e.g., credential_stuffing)')

    # List templates command
    list_parser = subparsers.add_parser('list', help='List available templates')

    # Create from template command
    create_parser = subparsers.add_parser('create', help='Create PIR from template')
    create_parser.add_argument('template', help='Template name')
    create_parser.add_argument('output', help='Output file path')
    create_parser.add_argument('--ticket', required=True, help='Ticket number')
    create_parser.add_argument('--customer', required=True, help='Customer name')
    create_parser.add_argument('--incident-type', help='Incident type')
    create_parser.add_argument('--severity', default='SEV1', help='Severity (SEV1-4)')

    args = parser.parse_args()

    manager = PIRTemplateManager()

    if args.command == 'save':
        metadata = {}
        if args.incident_type:
            metadata['incident_type'] = args.incident_type

        manager.save_as_template(
            source_docx=args.source,
            template_name=args.name,
            description=args.description or "",
            metadata=metadata
        )

    elif args.command == 'list':
        templates = manager.list_templates()

        if not templates:
            print("No templates found.")
        else:
            print(f"\n{'=' * 80}")
            print("AVAILABLE PIR TEMPLATES")
            print(f"{'=' * 80}\n")

            for template in templates:
                print(f"Name: {template['template_name']}")
                print(f"Description: {template.get('description', 'N/A')}")
                print(f"Incident Type: {template.get('incident_type', 'generic')}")
                print(f"Created: {template['created_date']}")
                print(f"Sections: {len(template.get('sections', []))}")
                print()

    elif args.command == 'create':
        incident_data = {
            'ticket_number': args.ticket,
            'customer_name': args.customer,
            'incident_type': args.incident_type or 'Security Incident',
            'severity': args.severity,
            'analyst_name': 'Security Specialist Agent (Maia)'
        }

        manager.create_from_template(
            template_name=args.template,
            output_path=args.output,
            incident_data=incident_data
        )

    else:
        parser.print_help()


if __name__ == '__main__':
    main()
