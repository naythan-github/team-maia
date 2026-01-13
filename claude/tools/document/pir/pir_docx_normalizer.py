#!/usr/bin/env python3
"""
PIR Document Normalizer - Fixes formatting inconsistencies in Word documents.

Primary use: Post-process PIR (Post-Incident Review) documents to ensure
consistent table widths, cell padding, and styling.

Usage:
    python3 pir_docx_normalizer.py <input.docx> [output.docx]
    python3 pir_docx_normalizer.py --fix-template <template.docx>
"""

import argparse
import sys
from pathlib import Path

# Optional dependency - graceful degradation for importability
try:
    from docx import Document
    from docx.shared import Pt, Inches, Twips
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    Pt = Inches = Twips = None
    qn = OxmlElement = None


def _check_docx_available():
    """Raise ImportError if python-docx is not installed."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx required. Install with: pip install python-docx")


def set_table_full_width(table):
    """
    Set table to 100% page width.

    IMPORTANT: OOXML element ordering is critical. Elements must appear in this order:
    tblStyle → tblpPr → tblW → jc → ... → tblLayout → tblCellMar → tblLook

    Word ignores tblW if it appears before tblStyle.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr

    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove existing width element if present
    existing_width = tblPr.find(qn('w:tblW'))
    if existing_width is not None:
        tblPr.remove(existing_width)

    # Create new width element set to 100%
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), '5000')  # 5000 = 100% in Word's percentage units

    # Insert tblW AFTER tblStyle (if present) - OOXML ordering requirement
    tblStyle = tblPr.find(qn('w:tblStyle'))
    if tblStyle is not None:
        tblStyle.addnext(tblW)
    else:
        tblPr.insert(0, tblW)

    # Disable autofit to respect our width setting
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')


def get_original_cell_widths(table):
    """Extract original cell widths from first row, preserving ratios."""
    widths = []
    if not table.rows:
        return widths

    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.tcPr
        width = 1440  # Default: 1 inch in twips if not specified

        if tcPr is not None:
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is not None:
                w_val = tcW.get(qn('w:w'))
                w_type = tcW.get(qn('w:type'))
                if w_val:
                    if w_type == 'pct':
                        # Already percentage - use directly
                        width = int(w_val)
                    else:
                        # dxa (twips) or other - convert to numeric
                        width = int(w_val)
        widths.append(width)

    return widths


def convert_widths_to_percentages(widths):
    """Convert absolute widths to percentage values (5000 = 100%)."""
    total = sum(widths)
    if total == 0:
        # Fallback to equal distribution
        num_cols = len(widths)
        return [5000 // num_cols] * num_cols

    # Convert each width to percentage of total, scaled to Word's 5000 = 100%
    percentages = []
    for w in widths:
        pct = int((w / total) * 5000)
        percentages.append(pct)

    # Ensure they sum to exactly 5000 (handle rounding)
    diff = 5000 - sum(percentages)
    if diff != 0 and percentages:
        percentages[-1] += diff

    return percentages


def set_cell_widths_preserve_ratios(table):
    """Set cell widths to percentages while PRESERVING original ratios."""
    if not table.rows:
        return

    # Get original widths and convert to percentages
    original_widths = get_original_cell_widths(table)
    pct_widths = convert_widths_to_percentages(original_widths)

    # Apply percentage widths to all rows
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i >= len(pct_widths):
                continue  # Handle merged cells

            tc = cell._tc
            tcPr = tc.tcPr

            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)

            # Remove existing width
            existing_width = tcPr.find(qn('w:tcW'))
            if existing_width is not None:
                tcPr.remove(existing_width)

            # Set percentage width preserving ratio
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:type'), 'pct')
            tcW.set(qn('w:w'), str(pct_widths[i]))
            tcPr.insert(0, tcW)


def calculate_content_aware_widths(table):
    """
    Calculate column widths based on actual content in cells.

    Analyzes text length across all rows to determine optimal proportions.
    Applies min/max constraints to prevent extreme column sizes.

    Returns list of widths in Word percentage units (sum = 5000 = 100%).
    """
    if not table.rows:
        return []

    num_cols = len(table.columns)
    col_max_lengths = [0] * num_cols

    # Find max content length per column (across all rows)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < num_cols:
                # Use max line length (handles multi-line cells)
                text = cell.text or ""
                lines = text.split('\n')
                max_line = max((len(line) for line in lines), default=0)
                col_max_lengths[i] = max(col_max_lengths[i], max_line)

    # Handle empty table
    total_chars = sum(col_max_lengths)
    if total_chars == 0:
        # Equal distribution fallback
        base_width = 5000 // num_cols
        widths = [base_width] * num_cols
        widths[-1] += 5000 - sum(widths)
        return widths

    # Convert character lengths to initial percentages
    raw_widths = []
    for length in col_max_lengths:
        pct = int((length / total_chars) * 5000)
        raw_widths.append(pct)

    # Apply constraints: min 10% (500), max 60% (3000) per column
    MIN_PCT = 500   # 10% minimum
    MAX_PCT = 3000  # 60% maximum

    constrained = []
    for w in raw_widths:
        constrained.append(max(MIN_PCT, min(MAX_PCT, w)))

    # Normalize to sum to exactly 5000
    current_total = sum(constrained)
    if current_total == 0:
        current_total = 1  # Prevent division by zero

    normalized = [int(w * 5000 / current_total) for w in constrained]

    # Fix rounding errors - adjust last column
    diff = 5000 - sum(normalized)
    if normalized:
        normalized[-1] += diff

    return normalized


def set_cell_widths_content_aware(table):
    """Set cell widths based on content analysis."""
    if not table.rows:
        return

    pct_widths = calculate_content_aware_widths(table)

    # Apply calculated widths to all rows
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i >= len(pct_widths):
                continue  # Handle merged cells

            tc = cell._tc
            tcPr = tc.tcPr

            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)

            # Remove existing width
            existing_width = tcPr.find(qn('w:tcW'))
            if existing_width is not None:
                tcPr.remove(existing_width)

            # Set content-aware width
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:type'), 'pct')
            tcW.set(qn('w:w'), str(pct_widths[i]))
            tcPr.insert(0, tcW)


def apply_table_style(table, doc, style_name):
    """
    Apply a named table style to a table.

    Args:
        table: The table to style
        doc: The Document object (needed to access styles)
        style_name: Name of the style to apply (e.g., "_Orro Table 1")

    Returns True if style applied, False if style not found.
    """
    try:
        # Check if style exists in document
        style = doc.styles[style_name]
        if style.type != 3:  # 3 = TABLE style
            return False
        table.style = style_name
        return True
    except KeyError:
        return False


def enable_first_row_header(table):
    """Enable first row as header (for table styles that format headers differently)."""
    tbl = table._tbl
    tblPr = tbl.tblPr

    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Set tblLook to enable first row formatting
    tblLook = tblPr.find(qn('w:tblLook'))
    if tblLook is None:
        tblLook = OxmlElement('w:tblLook')
        tblPr.append(tblLook)

    tblLook.set(qn('w:firstRow'), '1')
    tblLook.set(qn('w:lastRow'), '0')
    tblLook.set(qn('w:firstColumn'), '0')
    tblLook.set(qn('w:lastColumn'), '0')
    tblLook.set(qn('w:noHBand'), '1')
    tblLook.set(qn('w:noVBand'), '1')


def update_table_grid(table, pct_widths, page_width_twips=9360):
    """
    Update tblGrid column widths to match page width.

    The tblGrid stores absolute column widths in twips (dxa).
    These must be updated to reflect the actual page width for
    Word to render the table correctly at 100% width.

    Args:
        table: The table to update
        pct_widths: List of percentage widths (sum = 5000 = 100%)
        page_width_twips: Page text width in twips (default: 9360 = 6.5" for letter)
    """
    tbl = table._tbl

    # Find or create tblGrid
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        # Insert after tblPr
        tblPr = tbl.tblPr
        if tblPr is not None:
            tblPr.addnext(tblGrid)
        else:
            tbl.insert(0, tblGrid)
    else:
        # Clear existing gridCol elements
        for gridCol in list(tblGrid.findall(qn('w:gridCol'))):
            tblGrid.remove(gridCol)

    # Calculate absolute widths from percentages
    for pct in pct_widths:
        gridCol = OxmlElement('w:gridCol')
        # Convert percentage (5000 = 100%) to twips
        width_twips = int((pct / 5000) * page_width_twips)
        gridCol.set(qn('w:w'), str(width_twips))
        tblGrid.append(gridCol)


def apply_orro_borders(table):
    """
    Apply Orro corporate border styling with explicit RGB colors.

    Orro style borders:
    - Bottom: Purple #7030A0, 8pt (1pt = 8 eighths)
    - Inside horizontal: Light purple #CBC6F3, 4pt
    - No vertical borders (clean Orro look)

    Uses explicit RGB colors instead of theme colors to ensure
    consistent rendering regardless of document theme.

    Also applies explicit borders to first row cells to override
    the style's theme-dependent firstRow borders.
    """
    ORRO_PURPLE = "7030A0"
    ORRO_LIGHT_PURPLE = "CBC6F3"

    tbl = table._tbl
    tblPr = tbl.tblPr

    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove existing borders
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)

    # Create new borders element
    tblBorders = OxmlElement('w:tblBorders')

    # Bottom border - Orro purple, 8pt (64 eighths)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')  # 1pt in eighths
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), ORRO_PURPLE)
    tblBorders.append(bottom)

    # Inside horizontal - light purple, 4pt
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '4')
    insideH.set(qn('w:space'), '0')
    insideH.set(qn('w:color'), ORRO_LIGHT_PURPLE)
    tblBorders.append(insideH)

    # Explicitly set no vertical borders
    for border_name in ['left', 'right', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)

    # Top border - nil (no top border)
    top_border = OxmlElement('w:top')
    top_border.set(qn('w:val'), 'nil')
    tblBorders.append(top_border)

    # Insert borders after tblW (maintain OOXML order)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None:
        tblW.addnext(tblBorders)
    else:
        tblPr.insert(0, tblBorders)

    # Apply explicit borders to first row cells to override style's theme-dependent borders
    if table.rows:
        for cell in table.rows[0].cells:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)

            # Remove existing cell borders
            existing_tc_borders = tcPr.find(qn('w:tcBorders'))
            if existing_tc_borders is not None:
                tcPr.remove(existing_tc_borders)

            # Create cell borders with explicit Orro purple (no themeColor)
            tcBorders = OxmlElement('w:tcBorders')

            # Bottom border for header row - Orro purple
            tc_bottom = OxmlElement('w:bottom')
            tc_bottom.set(qn('w:val'), 'single')
            tc_bottom.set(qn('w:sz'), '8')
            tc_bottom.set(qn('w:space'), '0')
            tc_bottom.set(qn('w:color'), ORRO_PURPLE)
            tcBorders.append(tc_bottom)

            # No other borders for cells
            for border_name in ['top', 'left', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)

            tcPr.append(tcBorders)


def apply_table_font(table, font_name="Aptos", font_size_pt=10):
    """
    Apply consistent font styling to all text in a table.

    Args:
        table: The table to style
        font_name: Font family name (default: Aptos - Orro corporate font)
        font_size_pt: Font size in points (default: 10)
    """
    from docx.shared import Pt

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size_pt)


def normalize_paragraph_spacing(doc):
    """
    Normalize paragraph spacing throughout the document.

    Applies consistent spacing based on paragraph style:
    - Body text (Normal): 6pt after
    - Headings: 12pt before, 6pt after
    - List items: 2pt before/after (tight)

    Table cell spacing is handled separately in normalize_table_cell_spacing().
    """
    BODY_SPACE_AFTER = Pt(6)
    HEADING_SPACE_BEFORE = Pt(12)
    HEADING_SPACE_AFTER = Pt(6)
    LIST_SPACE = Pt(2)

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""

        if 'Heading' in style_name:
            # Headings: 12pt before, 6pt after
            para.paragraph_format.space_before = HEADING_SPACE_BEFORE
            para.paragraph_format.space_after = HEADING_SPACE_AFTER
        elif 'List' in style_name:
            # List items: tight 2pt spacing
            para.paragraph_format.space_before = LIST_SPACE
            para.paragraph_format.space_after = LIST_SPACE
        elif style_name == 'Normal' or style_name == '':
            # Body text: 6pt after (no explicit before)
            para.paragraph_format.space_after = BODY_SPACE_AFTER


def normalize_table_cell_spacing(table):
    """
    Apply minimal spacing to table cell paragraphs.

    Table cells should have tight 2pt before/after spacing
    to keep rows compact.
    """
    TABLE_CELL_SPACE = Pt(2)

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = TABLE_CELL_SPACE
                para.paragraph_format.space_after = TABLE_CELL_SPACE


def set_cell_margins(table, top=50, bottom=50, left=100, right=100):
    """Set consistent cell margins (in twips: 1440 twips = 1 inch)."""
    tbl = table._tbl
    tblPr = tbl.tblPr

    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove existing margins
    existing_margins = tblPr.find(qn('w:tblCellMar'))
    if existing_margins is not None:
        tblPr.remove(existing_margins)

    # Create cell margins element
    tblCellMar = OxmlElement('w:tblCellMar')

    for side, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin = OxmlElement(f'w:{side}')
        margin.set(qn('w:w'), str(value))
        margin.set(qn('w:type'), 'dxa')
        tblCellMar.append(margin)

    tblPr.append(tblCellMar)


def normalize_document(doc_path: Path, output_path: Path = None, verbose: bool = True,
                       content_aware: bool = True, table_style: str = None) -> dict:
    """
    Normalize a Word document's formatting.

    Args:
        doc_path: Input document path
        output_path: Output path (default: overwrite input)
        verbose: Print progress
        content_aware: Use content-aware column sizing (default: True)
                      If False, preserves existing width ratios
        table_style: Name of table style to apply (e.g., "_Orro Table 1")

    Returns dict with statistics about changes made.
    """
    doc = Document(str(doc_path))

    stats = {
        "tables_fixed": 0,
        "tables_styled": 0,
        "total_tables": len(doc.tables),
        "input": str(doc_path),
        "output": str(output_path or doc_path),
        "mode": "content-aware" if content_aware else "preserve-ratios",
        "table_style": table_style
    }

    for i, table in enumerate(doc.tables):
        # Apply table style if specified
        if table_style:
            if apply_table_style(table, doc, table_style):
                enable_first_row_header(table)
                stats["tables_styled"] += 1

        # Fix table width to 100%
        set_table_full_width(table)

        # Calculate and set column widths
        if content_aware:
            pct_widths = calculate_content_aware_widths(table)
        else:
            original_widths = get_original_cell_widths(table)
            pct_widths = convert_widths_to_percentages(original_widths)

        # Apply percentage widths to cells
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                if j >= len(pct_widths):
                    continue
                tc = cell._tc
                tcPr = tc.tcPr
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)
                existing_width = tcPr.find(qn('w:tcW'))
                if existing_width is not None:
                    tcPr.remove(existing_width)
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:type'), 'pct')
                tcW.set(qn('w:w'), str(pct_widths[j]))
                tcPr.insert(0, tcW)

        # Update tblGrid to match page width (CRITICAL for Word rendering)
        update_table_grid(table, pct_widths)

        # Apply Orro borders with explicit RGB colors (overrides theme-dependent style borders)
        apply_orro_borders(table)

        # Apply Aptos 8pt font to all table text
        apply_table_font(table, font_name="Aptos", font_size_pt=8)

        # Apply minimal spacing to table cell paragraphs
        normalize_table_cell_spacing(table)

        # Set consistent cell margins (skip if style applied - style has its own margins)
        if not table_style:
            set_cell_margins(table)

        stats["tables_fixed"] += 1

        if verbose:
            style_info = f" [styled]" if table_style and stats["tables_styled"] > i else ""
            print(f"  Fixed table {i + 1}/{stats['total_tables']}: {len(table.rows)}x{len(table.columns)}{style_info}")

    # Normalize paragraph spacing for all document paragraphs
    normalize_paragraph_spacing(doc)

    if verbose:
        print(f"  Normalized paragraph spacing")

    # Save to output path (or overwrite input if not specified)
    save_path = output_path or doc_path
    doc.save(str(save_path))

    if verbose:
        print(f"\n✅ Normalized {stats['tables_fixed']} tables → {save_path}")

    return stats


def main():
    parser = argparse.ArgumentParser(
        description="Normalize PIR document formatting (tables to 100% width, etc.)"
    )
    parser.add_argument("input", help="Input .docx file")
    parser.add_argument("output", nargs="?", help="Output .docx file (default: overwrite input)")
    parser.add_argument("--fix-template", action="store_true",
                        help="Fix a template file in-place")
    parser.add_argument("--preserve-ratios", action="store_true",
                        help="Preserve existing column width ratios instead of content-aware sizing")
    parser.add_argument("--table-style", type=str, default=None,
                        help="Table style to apply (e.g., '_Orro Table 1')")
    parser.add_argument("-q", "--quiet", action="store_true",
                        help="Suppress verbose output")

    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"ERROR: File not found: {input_path}")
        sys.exit(1)

    if not input_path.suffix.lower() == '.docx':
        print(f"ERROR: Not a .docx file: {input_path}")
        sys.exit(1)

    output_path = Path(args.output) if args.output else None

    if args.fix_template:
        print(f"Fixing template: {input_path}")
        output_path = input_path  # Overwrite template

    stats = normalize_document(input_path, output_path, verbose=not args.quiet,
                               content_aware=not args.preserve_ratios,
                               table_style=args.table_style)

    return 0 if stats["tables_fixed"] > 0 else 1


if __name__ == "__main__":
    sys.exit(main())
