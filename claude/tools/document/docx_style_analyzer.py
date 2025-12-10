#!/usr/bin/env python3
"""
Word Document Style Analyzer - Identifies used vs unused styles.

Analyzes .docx files to determine:
- All styles defined in the document/template
- Which styles are actually used in content
- Usage counts per style
- Recommendations for cleanup

Usage:
    python3 docx_style_analyzer.py <document.docx>
    python3 docx_style_analyzer.py <document.docx> --json
    python3 docx_style_analyzer.py <document.docx> --verbose
"""

import argparse
import json
import sys
from collections import defaultdict
from pathlib import Path
from typing import Dict, List, Set, Tuple

try:
    from docx import Document
    from docx.styles.style import BaseStyle
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# Style type mapping for readable output
STYLE_TYPE_NAMES = {
    WD_STYLE_TYPE.PARAGRAPH: "Paragraph",
    WD_STYLE_TYPE.CHARACTER: "Character",
    WD_STYLE_TYPE.TABLE: "Table",
    WD_STYLE_TYPE.LIST: "List",
}


def get_all_defined_styles(doc: Document) -> Dict[str, dict]:
    """
    Extract all styles defined in the document.

    Returns dict mapping style_name -> {type, builtin, hidden, base_style}
    """
    styles = {}

    for style in doc.styles:
        try:
            style_info = {
                "name": style.name,
                "type": STYLE_TYPE_NAMES.get(style.type, f"Unknown({style.type})"),
                "type_id": style.type,
                "builtin": style.builtin,
                "hidden": style.hidden if hasattr(style, 'hidden') else False,
                "base_style": style.base_style.name if style.base_style else None,
            }
            styles[style.name] = style_info
        except Exception as e:
            # Some styles may have issues - skip them
            continue

    return styles


def analyze_paragraph_styles(doc: Document) -> Dict[str, int]:
    """
    Count usage of paragraph styles in document body.

    Returns dict mapping style_name -> usage_count
    """
    usage = defaultdict(int)

    # Document body paragraphs
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        usage[style_name] += 1

    # Table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    style_name = para.style.name if para.style else "Normal"
                    usage[style_name] += 1

    # Header/footer paragraphs
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.is_linked_to_previous is False:
                for para in header.paragraphs:
                    style_name = para.style.name if para.style else "Normal"
                    usage[style_name] += 1

        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                for para in footer.paragraphs:
                    style_name = para.style.name if para.style else "Normal"
                    usage[style_name] += 1

    return dict(usage)


def analyze_character_styles(doc: Document) -> Dict[str, int]:
    """
    Count usage of character (run) styles.

    Returns dict mapping style_name -> usage_count
    """
    usage = defaultdict(int)

    def process_paragraphs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                if run.style and run.style.name:
                    usage[run.style.name] += 1

    # Document body
    process_paragraphs(doc.paragraphs)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    # Headers/footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.is_linked_to_previous is False:
                process_paragraphs(header.paragraphs)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                process_paragraphs(footer.paragraphs)

    return dict(usage)


def analyze_table_styles(doc: Document) -> Dict[str, int]:
    """
    Count usage of table styles.

    Returns dict mapping style_name -> usage_count
    """
    usage = defaultdict(int)

    for table in doc.tables:
        style_name = table.style.name if table.style else "Table Grid"
        usage[style_name] += 1

    return dict(usage)


def analyze_document(doc_path: Path, verbose: bool = False) -> dict:
    """
    Comprehensive style analysis of a Word document.

    Returns analysis results including:
    - defined_styles: All styles in document
    - used_styles: Styles actually referenced in content
    - unused_styles: Defined but not used
    - usage_counts: How many times each style is used
    - recommendations: Cleanup suggestions
    """
    doc = Document(str(doc_path))

    # Get all defined styles
    defined_styles = get_all_defined_styles(doc)

    # Analyze usage by type
    para_usage = analyze_paragraph_styles(doc)
    char_usage = analyze_character_styles(doc)
    table_usage = analyze_table_styles(doc)

    # Combine all usage
    all_usage = {}
    all_usage.update({k: {"count": v, "type": "Paragraph"} for k, v in para_usage.items()})
    all_usage.update({k: {"count": v, "type": "Character"} for k, v in char_usage.items()})
    all_usage.update({k: {"count": v, "type": "Table"} for k, v in table_usage.items()})

    # Identify used and unused styles
    used_style_names = set(all_usage.keys())
    defined_style_names = set(defined_styles.keys())

    # Unused = defined but not in content
    unused_styles = []
    for name in defined_style_names:
        if name not in used_style_names:
            style_info = defined_styles[name]
            unused_styles.append({
                "name": name,
                "type": style_info["type"],
                "builtin": style_info["builtin"],
                "hidden": style_info["hidden"],
            })

    # Sort unused by type, then name
    unused_styles.sort(key=lambda x: (x["type"], x["name"]))

    # Used styles with counts
    used_styles = []
    for name, info in all_usage.items():
        style_defined = name in defined_styles
        used_styles.append({
            "name": name,
            "type": info["type"],
            "count": info["count"],
            "defined_in_doc": style_defined,
        })

    # Sort by count descending
    used_styles.sort(key=lambda x: -x["count"])

    # Build recommendations
    recommendations = []

    # Count removable styles (non-builtin, unused)
    removable = [s for s in unused_styles if not s["builtin"] and not s["hidden"]]
    if removable:
        recommendations.append(f"Can remove {len(removable)} custom unused styles")

    # Check for duplicate-looking styles
    style_bases = defaultdict(list)
    for name in defined_style_names:
        base = name.replace(" ", "").lower()
        style_bases[base].append(name)
    duplicates = {k: v for k, v in style_bases.items() if len(v) > 1}
    if duplicates:
        for base, names in duplicates.items():
            if len(names) > 1:
                recommendations.append(f"Potential duplicates: {', '.join(names)}")

    # Summary stats
    stats = {
        "file": str(doc_path),
        "total_defined": len(defined_styles),
        "total_used": len(used_styles),
        "total_unused": len(unused_styles),
        "paragraph_styles_used": len(para_usage),
        "character_styles_used": len(char_usage),
        "table_styles_used": len(table_usage),
        "builtin_unused": len([s for s in unused_styles if s["builtin"]]),
        "custom_unused": len([s for s in unused_styles if not s["builtin"]]),
    }

    return {
        "stats": stats,
        "defined_styles": defined_styles,
        "used_styles": used_styles,
        "unused_styles": unused_styles,
        "recommendations": recommendations,
    }


def print_report(analysis: dict, verbose: bool = False):
    """Print human-readable analysis report."""
    stats = analysis["stats"]

    print(f"\n{'='*60}")
    print(f"WORD DOCUMENT STYLE ANALYSIS")
    print(f"{'='*60}")
    print(f"File: {stats['file']}")
    print()

    # Summary
    print(f"SUMMARY")
    print(f"-" * 40)
    print(f"  Styles defined:     {stats['total_defined']:>4}")
    print(f"  Styles USED:        {stats['total_used']:>4}")
    print(f"  Styles UNUSED:      {stats['total_unused']:>4}")
    print()
    print(f"  Breakdown by type:")
    print(f"    Paragraph styles: {stats['paragraph_styles_used']:>4} used")
    print(f"    Character styles: {stats['character_styles_used']:>4} used")
    print(f"    Table styles:     {stats['table_styles_used']:>4} used")
    print()

    # Used styles with counts
    print(f"STYLES IN USE")
    print(f"-" * 40)
    for style in analysis["used_styles"]:
        defined = "✓" if style["defined_in_doc"] else "?"
        print(f"  [{defined}] {style['name']:<30} ({style['count']:>3}x) [{style['type']}]")
    print()

    # Unused styles
    if analysis["unused_styles"]:
        print(f"UNUSED STYLES ({stats['total_unused']} total)")
        print(f"-" * 40)

        if verbose:
            for style in analysis["unused_styles"]:
                builtin_tag = "[builtin]" if style["builtin"] else "[custom]"
                hidden_tag = "[hidden]" if style["hidden"] else ""
                print(f"  {style['name']:<35} {builtin_tag} {hidden_tag}")
        else:
            # Show counts by type
            by_type = defaultdict(list)
            for s in analysis["unused_styles"]:
                by_type[s["type"]].append(s)

            for stype, styles in sorted(by_type.items()):
                custom = len([s for s in styles if not s["builtin"]])
                builtin = len([s for s in styles if s["builtin"]])
                print(f"  {stype}: {len(styles)} unused ({custom} custom, {builtin} builtin)")

            print(f"\n  (Use --verbose to see full list)")
    print()

    # Recommendations
    if analysis["recommendations"]:
        print(f"RECOMMENDATIONS")
        print(f"-" * 40)
        for rec in analysis["recommendations"]:
            print(f"  • {rec}")
        print()

    # Cleanup potential
    removable = stats["custom_unused"]
    if removable > 0:
        pct = (removable / stats["total_defined"]) * 100
        print(f"CLEANUP POTENTIAL")
        print(f"-" * 40)
        print(f"  {removable} custom styles can be safely removed ({pct:.0f}% of defined)")
        print()


def main():
    parser = argparse.ArgumentParser(
        description="Analyze Word document styles - identify used vs unused"
    )
    parser.add_argument("document", help="Path to .docx file to analyze")
    parser.add_argument("--json", action="store_true", help="Output as JSON")
    parser.add_argument("--verbose", "-v", action="store_true",
                        help="Show detailed unused style list")

    args = parser.parse_args()

    doc_path = Path(args.document)
    if not doc_path.exists():
        print(f"ERROR: File not found: {doc_path}")
        sys.exit(1)

    if not doc_path.suffix.lower() == '.docx':
        print(f"ERROR: Not a .docx file: {doc_path}")
        sys.exit(1)

    try:
        analysis = analyze_document(doc_path, verbose=args.verbose)

        if args.json:
            # JSON output (exclude full defined_styles for brevity)
            output = {
                "stats": analysis["stats"],
                "used_styles": analysis["used_styles"],
                "unused_count": len(analysis["unused_styles"]),
                "recommendations": analysis["recommendations"],
            }
            print(json.dumps(output, indent=2))
        else:
            print_report(analysis, verbose=args.verbose)

        return 0

    except Exception as e:
        print(f"ERROR: Failed to analyze document: {e}")
        if args.verbose:
            import traceback
            traceback.print_exc()
        return 1


if __name__ == "__main__":
    sys.exit(main())
