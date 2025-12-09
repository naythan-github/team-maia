#!/usr/bin/env python3
"""
PIR Document Normalizer - Fixes formatting inconsistencies in Word documents.

Primary use: Post-process PIR (Post-Incident Review) documents to ensure
consistent table widths, cell padding, and styling.

Usage:
    python3 pir_docx_normalizer.py <input.docx> [output.docx]
    python3 pir_docx_normalizer.py --fix-template <template.docx>
"""

import argparse
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Twips
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def set_table_full_width(table):
    """Set table to 100% page width."""
    tbl = table._tbl
    tblPr = tbl.tblPr

    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove existing width element if present
    existing_width = tblPr.find(qn('w:tblW'))
    if existing_width is not None:
        tblPr.remove(existing_width)

    # Create new width element set to 100%
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), '5000')  # 5000 = 100% in Word's percentage units
    tblPr.insert(0, tblW)

    # Disable autofit to respect our width setting
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')


def get_original_cell_widths(table):
    """Extract original cell widths from first row, preserving ratios."""
    widths = []
    if not table.rows:
        return widths

    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.tcPr
        width = 1440  # Default: 1 inch in twips if not specified

        if tcPr is not None:
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is not None:
                w_val = tcW.get(qn('w:w'))
                w_type = tcW.get(qn('w:type'))
                if w_val:
                    if w_type == 'pct':
                        # Already percentage - use directly
                        width = int(w_val)
                    else:
                        # dxa (twips) or other - convert to numeric
                        width = int(w_val)
        widths.append(width)

    return widths


def convert_widths_to_percentages(widths):
    """Convert absolute widths to percentage values (5000 = 100%)."""
    total = sum(widths)
    if total == 0:
        # Fallback to equal distribution
        num_cols = len(widths)
        return [5000 // num_cols] * num_cols

    # Convert each width to percentage of total, scaled to Word's 5000 = 100%
    percentages = []
    for w in widths:
        pct = int((w / total) * 5000)
        percentages.append(pct)

    # Ensure they sum to exactly 5000 (handle rounding)
    diff = 5000 - sum(percentages)
    if diff != 0 and percentages:
        percentages[-1] += diff

    return percentages


def set_cell_widths_preserve_ratios(table):
    """Set cell widths to percentages while PRESERVING original ratios."""
    if not table.rows:
        return

    # Get original widths and convert to percentages
    original_widths = get_original_cell_widths(table)
    pct_widths = convert_widths_to_percentages(original_widths)

    # Apply percentage widths to all rows
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i >= len(pct_widths):
                continue  # Handle merged cells

            tc = cell._tc
            tcPr = tc.tcPr

            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)

            # Remove existing width
            existing_width = tcPr.find(qn('w:tcW'))
            if existing_width is not None:
                tcPr.remove(existing_width)

            # Set percentage width preserving ratio
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:type'), 'pct')
            tcW.set(qn('w:w'), str(pct_widths[i]))
            tcPr.insert(0, tcW)


def calculate_content_aware_widths(table):
    """
    Calculate column widths based on actual content in cells.

    Analyzes text length across all rows to determine optimal proportions.
    Applies min/max constraints to prevent extreme column sizes.

    Returns list of widths in Word percentage units (sum = 5000 = 100%).
    """
    if not table.rows:
        return []

    num_cols = len(table.columns)
    col_max_lengths = [0] * num_cols

    # Find max content length per column (across all rows)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < num_cols:
                # Use max line length (handles multi-line cells)
                text = cell.text or ""
                lines = text.split('\n')
                max_line = max((len(line) for line in lines), default=0)
                col_max_lengths[i] = max(col_max_lengths[i], max_line)

    # Handle empty table
    total_chars = sum(col_max_lengths)
    if total_chars == 0:
        # Equal distribution fallback
        base_width = 5000 // num_cols
        widths = [base_width] * num_cols
        widths[-1] += 5000 - sum(widths)
        return widths

    # Convert character lengths to initial percentages
    raw_widths = []
    for length in col_max_lengths:
        pct = int((length / total_chars) * 5000)
        raw_widths.append(pct)

    # Apply constraints: min 10% (500), max 60% (3000) per column
    MIN_PCT = 500   # 10% minimum
    MAX_PCT = 3000  # 60% maximum

    constrained = []
    for w in raw_widths:
        constrained.append(max(MIN_PCT, min(MAX_PCT, w)))

    # Normalize to sum to exactly 5000
    current_total = sum(constrained)
    if current_total == 0:
        current_total = 1  # Prevent division by zero

    normalized = [int(w * 5000 / current_total) for w in constrained]

    # Fix rounding errors - adjust last column
    diff = 5000 - sum(normalized)
    if normalized:
        normalized[-1] += diff

    return normalized


def set_cell_widths_content_aware(table):
    """Set cell widths based on content analysis."""
    if not table.rows:
        return

    pct_widths = calculate_content_aware_widths(table)

    # Apply calculated widths to all rows
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i >= len(pct_widths):
                continue  # Handle merged cells

            tc = cell._tc
            tcPr = tc.tcPr

            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)

            # Remove existing width
            existing_width = tcPr.find(qn('w:tcW'))
            if existing_width is not None:
                tcPr.remove(existing_width)

            # Set content-aware width
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:type'), 'pct')
            tcW.set(qn('w:w'), str(pct_widths[i]))
            tcPr.insert(0, tcW)


def set_cell_margins(table, top=50, bottom=50, left=100, right=100):
    """Set consistent cell margins (in twips: 1440 twips = 1 inch)."""
    tbl = table._tbl
    tblPr = tbl.tblPr

    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove existing margins
    existing_margins = tblPr.find(qn('w:tblCellMar'))
    if existing_margins is not None:
        tblPr.remove(existing_margins)

    # Create cell margins element
    tblCellMar = OxmlElement('w:tblCellMar')

    for side, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin = OxmlElement(f'w:{side}')
        margin.set(qn('w:w'), str(value))
        margin.set(qn('w:type'), 'dxa')
        tblCellMar.append(margin)

    tblPr.append(tblCellMar)


def normalize_document(doc_path: Path, output_path: Path = None, verbose: bool = True,
                       content_aware: bool = True) -> dict:
    """
    Normalize a Word document's formatting.

    Args:
        doc_path: Input document path
        output_path: Output path (default: overwrite input)
        verbose: Print progress
        content_aware: Use content-aware column sizing (default: True)
                      If False, preserves existing width ratios

    Returns dict with statistics about changes made.
    """
    doc = Document(str(doc_path))

    stats = {
        "tables_fixed": 0,
        "total_tables": len(doc.tables),
        "input": str(doc_path),
        "output": str(output_path or doc_path),
        "mode": "content-aware" if content_aware else "preserve-ratios"
    }

    for i, table in enumerate(doc.tables):
        # Fix table width to 100%
        set_table_full_width(table)

        # Set column widths
        if content_aware:
            set_cell_widths_content_aware(table)
        else:
            set_cell_widths_preserve_ratios(table)

        # Set consistent cell margins
        set_cell_margins(table)

        stats["tables_fixed"] += 1

        if verbose:
            print(f"  Fixed table {i + 1}/{stats['total_tables']}: {len(table.rows)}x{len(table.columns)}")

    # Save to output path (or overwrite input if not specified)
    save_path = output_path or doc_path
    doc.save(str(save_path))

    if verbose:
        print(f"\n✅ Normalized {stats['tables_fixed']} tables → {save_path}")

    return stats


def main():
    parser = argparse.ArgumentParser(
        description="Normalize PIR document formatting (tables to 100% width, etc.)"
    )
    parser.add_argument("input", help="Input .docx file")
    parser.add_argument("output", nargs="?", help="Output .docx file (default: overwrite input)")
    parser.add_argument("--fix-template", action="store_true",
                        help="Fix a template file in-place")
    parser.add_argument("--preserve-ratios", action="store_true",
                        help="Preserve existing column width ratios instead of content-aware sizing")
    parser.add_argument("-q", "--quiet", action="store_true",
                        help="Suppress verbose output")

    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"ERROR: File not found: {input_path}")
        sys.exit(1)

    if not input_path.suffix.lower() == '.docx':
        print(f"ERROR: Not a .docx file: {input_path}")
        sys.exit(1)

    output_path = Path(args.output) if args.output else None

    if args.fix_template:
        print(f"Fixing template: {input_path}")
        output_path = input_path  # Overwrite template

    stats = normalize_document(input_path, output_path, verbose=not args.quiet,
                               content_aware=not args.preserve_ratios)

    return 0 if stats["tables_fixed"] > 0 else 1


if __name__ == "__main__":
    sys.exit(main())
