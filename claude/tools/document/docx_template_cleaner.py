#!/usr/bin/env python3
"""
Word Document Template Cleaner - Creates minimal templates with only used styles.

Uses hybrid approach:
1. Analyze source document for used styles
2. Build style dependency chain (base styles)
3. Extract only required style definitions from XML
4. Create clean document with minimal styles

Usage:
    python3 docx_template_cleaner.py <input.docx> <output.docx>
    python3 docx_template_cleaner.py <input.docx> <output.docx> --dry-run
"""

import argparse
import json
import shutil
import sys
import tempfile
import zipfile
from collections import defaultdict
from copy import deepcopy
from pathlib import Path
from typing import Dict, List, Set
from xml.etree import ElementTree as ET

# Optional dependency - graceful degradation for importability
try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    WD_STYLE_TYPE = None


def _check_docx_available():
    """Raise ImportError if python-docx is not installed."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx required. Install with: pip install python-docx")


# XML namespaces used in .docx files
NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'w15': 'http://schemas.microsoft.com/office/word/2012/wordml',
}

# Register namespaces to preserve prefixes on output
for prefix, uri in NAMESPACES.items():
    ET.register_namespace(prefix, uri)


def get_used_styles_from_document(doc_path: Path) -> Dict[str, Set[str]]:
    """
    Extract all styles actually used in a document.

    Returns dict with keys: 'paragraph', 'character', 'table', 'numbering'
    Each value is a set of style names/IDs used.
    """
    doc = Document(str(doc_path))

    used = {
        'paragraph': set(),
        'character': set(),
        'table': set(),
    }

    def process_paragraphs(paragraphs):
        for para in paragraphs:
            if para.style:
                used['paragraph'].add(para.style.name)
                # Also get the style_id for XML matching
                if hasattr(para.style, 'style_id') and para.style.style_id:
                    used['paragraph'].add(para.style.style_id)
            for run in para.runs:
                if run.style and run.style.name:
                    used['character'].add(run.style.name)
                    if hasattr(run.style, 'style_id') and run.style.style_id:
                        used['character'].add(run.style.style_id)

    # Document body
    process_paragraphs(doc.paragraphs)

    # Tables
    for table in doc.tables:
        if table.style:
            used['table'].add(table.style.name)
            if hasattr(table.style, 'style_id') and table.style.style_id:
                used['table'].add(table.style.style_id)
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    # Headers/footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                try:
                    process_paragraphs(header.paragraphs)
                    for table in header.tables:
                        if table.style:
                            used['table'].add(table.style.name)
                        for row in table.rows:
                            for cell in row.cells:
                                process_paragraphs(cell.paragraphs)
                except (KeyError, AttributeError, TypeError):
                    pass

        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                try:
                    process_paragraphs(footer.paragraphs)
                    for table in footer.tables:
                        if table.style:
                            used['table'].add(table.style.name)
                        for row in table.rows:
                            for cell in row.cells:
                                process_paragraphs(cell.paragraphs)
                except (KeyError, AttributeError, TypeError):
                    pass

    return used


def parse_styles_xml(styles_xml_content: bytes) -> ET.Element:
    """Parse styles.xml content into ElementTree."""
    return ET.fromstring(styles_xml_content)


def get_style_dependencies(styles_root: ET.Element, used_style_ids: Set[str]) -> Set[str]:
    """
    Build complete set of style IDs including base style dependencies.

    Follows the basedOn chain to include parent styles.
    """
    w = '{' + NAMESPACES['w'] + '}'

    # Build map of styleId -> basedOn styleId
    base_map = {}
    style_names_to_ids = {}

    for style in styles_root.findall(f'.//{w}style'):
        style_id = style.get(f'{w}styleId')
        if not style_id:
            continue

        # Map name to ID
        name_elem = style.find(f'{w}name')
        if name_elem is not None:
            name = name_elem.get(f'{w}val')
            if name:
                style_names_to_ids[name] = style_id

        # Get basedOn
        based_on = style.find(f'{w}basedOn')
        if based_on is not None:
            base_id = based_on.get(f'{w}val')
            if base_id:
                base_map[style_id] = base_id

    # Convert any names in used_style_ids to IDs
    normalized_ids = set()
    for item in used_style_ids:
        if item in style_names_to_ids:
            normalized_ids.add(style_names_to_ids[item])
        else:
            normalized_ids.add(item)

    # Expand with base styles
    all_needed = set(normalized_ids)
    to_check = list(normalized_ids)

    while to_check:
        current = to_check.pop()
        if current in base_map:
            base_id = base_map[current]
            if base_id not in all_needed:
                all_needed.add(base_id)
                to_check.append(base_id)

    # Also include linked styles (paragraph <-> character links)
    for style in styles_root.findall(f'.//{w}style'):
        style_id = style.get(f'{w}styleId')
        if style_id in all_needed:
            link = style.find(f'{w}link')
            if link is not None:
                linked_id = link.get(f'{w}val')
                if linked_id:
                    all_needed.add(linked_id)

    return all_needed


def get_essential_styles() -> Set[str]:
    """
    Return set of style IDs that must always be kept.

    These are required by Word for proper document functioning.
    """
    return {
        'Normal',           # Default paragraph style
        'DefaultParagraphFont',  # Default character style
        'TableNormal',      # Default table style
        'NoList',           # Default list style
        'Heading1', 'Heading2', 'Heading3',  # Common headings (often referenced)
    }


def clean_latent_styles(latent_styles: ET.Element, keep_style_names: Set[str]) -> ET.Element:
    """
    Clean latentStyles to hide all built-in styles except those in use.

    Sets defaults to hide everything, adds exceptions only for actually used styles.
    Filters out Char styles (they inherit visibility from parent paragraph style).
    """
    w = '{' + NAMESPACES['w'] + '}'

    # Create new latentStyles with restrictive defaults
    new_latent = ET.Element(f'{w}latentStyles')

    # Restrictive defaults: hide all by default, show only when used
    new_latent.set(f'{w}defLockedState', '0')
    new_latent.set(f'{w}defUIPriority', '99')
    new_latent.set(f'{w}defSemiHidden', '1')  # Hidden by default
    new_latent.set(f'{w}defUnhideWhenUsed', '1')  # Show when used
    new_latent.set(f'{w}defQFormat', '0')
    new_latent.set(f'{w}count', '376')

    # Filter visible styles:
    # 1. Remove Char styles (they inherit from parent, don't need exceptions)
    # 2. Only include styles actually used in document
    visible_styles = {
        name for name in keep_style_names
        if not name.endswith(' Char')
    }

    # Always include Normal (essential for Word)
    visible_styles.add('Normal')

    for name in sorted(visible_styles):
        exc = ET.Element(f'{w}lsdException')
        exc.set(f'{w}name', name)
        exc.set(f'{w}semiHidden', '0')
        exc.set(f'{w}uiPriority', '1')  # Low priority = appears at top
        exc.set(f'{w}unhideWhenUsed', '0')
        exc.set(f'{w}qFormat', '1')
        new_latent.append(exc)

    return new_latent


def filter_styles_xml(styles_root: ET.Element, keep_style_ids: Set[str]) -> ET.Element:
    """
    Create new styles.xml with only the specified styles.

    Also cleans latentStyles to hide unused built-in styles.
    """
    w = '{' + NAMESPACES['w'] + '}'

    # Build name set from IDs for latent style filtering
    keep_style_names = set()
    for style in styles_root.findall(f'.//{w}style'):
        style_id = style.get(f'{w}styleId')
        if style_id in keep_style_ids:
            name_elem = style.find(f'{w}name')
            if name_elem is not None:
                keep_style_names.add(name_elem.get(f'{w}val'))

    # Create new root
    new_root = ET.Element(styles_root.tag, styles_root.attrib)

    # Copy document defaults (fonts, etc.) - always needed
    doc_defaults = styles_root.find(f'{w}docDefaults')
    if doc_defaults is not None:
        new_root.append(deepcopy(doc_defaults))

    # Create cleaned latent styles (hide unused built-ins)
    latent_styles = styles_root.find(f'{w}latentStyles')
    if latent_styles is not None:
        new_latent = clean_latent_styles(latent_styles, keep_style_names)
        new_root.append(new_latent)

    # Copy only needed styles
    kept_count = 0
    removed_count = 0

    for style in styles_root.findall(f'{w}style'):
        style_id = style.get(f'{w}styleId')

        # Always keep if in our keep set
        if style_id in keep_style_ids:
            new_root.append(deepcopy(style))
            kept_count += 1
        else:
            removed_count += 1

    return new_root, kept_count, removed_count


def fix_style_pane_filter(settings_xml: bytes) -> bytes:
    """
    Fix the stylePaneFormatFilter in settings.xml to show only styles in use.

    By default, Word may show latent (built-in) styles even when hidden.
    This fixes the filter to show only:
    - stylesInUse: 1 (styles actually used in document)
    - latentStyles: 0 (don't show built-in styles not in use)

    Returns modified settings.xml as bytes.
    """
    w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    root = ET.fromstring(settings_xml)

    # Find or create stylePaneFormatFilter
    filter_el = root.find(f'.//{w}stylePaneFormatFilter')

    if filter_el is None:
        # Create the element if it doesn't exist
        filter_el = ET.Element(f'{w}stylePaneFormatFilter')
        root.append(filter_el)

    # Set the filter to show only styles in use
    # Key attributes:
    # - stylesInUse="1" - show styles used in document
    # - latentStyles="0" - don't show built-in latent styles
    # - customStyles="0" - we'll show them via stylesInUse
    # - allStyles="0" - don't show all styles
    filter_el.set(f'{w}val', '3F01')  # Hex value for stylesInUse=1
    filter_el.set(f'{w}allStyles', '0')
    filter_el.set(f'{w}customStyles', '0')
    filter_el.set(f'{w}latentStyles', '0')  # KEY: Don't show latent styles
    filter_el.set(f'{w}stylesInUse', '1')   # KEY: Show only styles in use
    filter_el.set(f'{w}headingStyles', '0')
    filter_el.set(f'{w}numberingStyles', '0')
    filter_el.set(f'{w}tableStyles', '0')
    filter_el.set(f'{w}directFormattingOnRuns', '0')
    filter_el.set(f'{w}directFormattingOnParagraphs', '0')
    filter_el.set(f'{w}directFormattingOnNumbering', '0')
    filter_el.set(f'{w}directFormattingOnTables', '0')
    filter_el.set(f'{w}clearFormatting', '1')
    filter_el.set(f'{w}top3HeadingStyles', '0')
    filter_el.set(f'{w}visibleStyles', '1')  # Show visible styles
    filter_el.set(f'{w}alternateStyleNames', '0')

    # Return as XML bytes
    xml_str = ET.tostring(root, encoding='unicode')
    xml_str = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + xml_str
    return xml_str.encode('utf-8')


def clean_template(input_path: Path, output_path: Path, dry_run: bool = False) -> dict:
    """
    Create a clean template with only used styles.

    Args:
        input_path: Source document
        output_path: Destination for cleaned document
        dry_run: If True, only analyze without creating output

    Returns dict with statistics.
    """
    # Step 1: Get used styles from document content
    used_styles = get_used_styles_from_document(input_path)
    all_used = set()
    for category in used_styles.values():
        all_used.update(category)

    # Step 2: Read styles.xml from the docx
    with zipfile.ZipFile(input_path, 'r') as zf:
        styles_xml = zf.read('word/styles.xml')

    styles_root = parse_styles_xml(styles_xml)

    # Step 3: Expand with dependencies and essentials
    essential = get_essential_styles()
    keep_ids = get_style_dependencies(styles_root, all_used)
    keep_ids.update(essential)

    # Step 4: Filter styles
    new_styles_root, kept, removed = filter_styles_xml(styles_root, keep_ids)

    stats = {
        'input': str(input_path),
        'output': str(output_path),
        'original_styles': kept + removed,
        'kept_styles': kept,
        'removed_styles': removed,
        'used_paragraph': len(used_styles['paragraph']),
        'used_character': len(used_styles['character']),
        'used_table': len(used_styles['table']),
        'dry_run': dry_run,
    }

    if dry_run:
        return stats

    # Step 5: Create cleaned document
    # Copy original to temp, modify styles.xml, save to output
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir) / 'temp.docx'
        shutil.copy2(input_path, tmp_path)

        # Rewrite the docx with new styles.xml and fixed settings.xml
        with zipfile.ZipFile(tmp_path, 'r') as zf_in:
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zf_out:
                for item in zf_in.namelist():
                    if item == 'word/styles.xml':
                        # Write our filtered styles
                        new_xml = ET.tostring(new_styles_root, encoding='unicode')
                        # Add XML declaration
                        new_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + new_xml
                        zf_out.writestr(item, new_xml.encode('utf-8'))
                    elif item == 'word/settings.xml':
                        # Fix the style pane filter to show only styles in use
                        original_settings = zf_in.read(item)
                        fixed_settings = fix_style_pane_filter(original_settings)
                        zf_out.writestr(item, fixed_settings)
                    else:
                        # Copy unchanged
                        zf_out.writestr(item, zf_in.read(item))

    return stats


def print_report(stats: dict):
    """Print cleaning report."""
    print(f"\n{'='*60}")
    print("TEMPLATE CLEANING REPORT")
    print(f"{'='*60}")
    print(f"Input:  {stats['input']}")
    print(f"Output: {stats['output']}")
    print()

    print("STYLE ANALYSIS")
    print("-" * 40)
    print(f"  Original styles:  {stats['original_styles']:>4}")
    print(f"  Kept styles:      {stats['kept_styles']:>4}")
    print(f"  Removed styles:   {stats['removed_styles']:>4}")
    print()

    reduction = (stats['removed_styles'] / stats['original_styles'] * 100) if stats['original_styles'] > 0 else 0
    print(f"  Reduction: {reduction:.0f}%")
    print()

    if stats['dry_run']:
        print("⚠️  DRY RUN - No file created")
    else:
        print(f"✅ Clean template created: {stats['output']}")
    print()


def main():
    parser = argparse.ArgumentParser(
        description="Create minimal Word template with only used styles"
    )
    parser.add_argument("input", help="Input .docx file")
    parser.add_argument("output", help="Output .docx file (cleaned)")
    parser.add_argument("--dry-run", action="store_true",
                        help="Analyze only, don't create output")
    parser.add_argument("--json", action="store_true",
                        help="Output stats as JSON")

    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)

    if not input_path.exists():
        print(f"ERROR: File not found: {input_path}")
        sys.exit(1)

    if not input_path.suffix.lower() == '.docx':
        print(f"ERROR: Not a .docx file: {input_path}")
        sys.exit(1)

    if input_path.resolve() == output_path.resolve():
        print("ERROR: Input and output must be different files")
        sys.exit(1)

    try:
        stats = clean_template(input_path, output_path, dry_run=args.dry_run)

        if args.json:
            print(json.dumps(stats, indent=2))
        else:
            print_report(stats)

        return 0

    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    sys.exit(main())
