#!/usr/bin/env python3
"""
Tests for Word Document Style Creator.

TDD tests for detecting formatting patterns and creating named styles.
"""

import tempfile
import unittest
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET
from dataclasses import dataclass

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

import sys
sys.path.insert(0, str(Path(__file__).parent.parent / "claude" / "tools" / "document"))


@unittest.skipUnless(DOCX_AVAILABLE, "python-docx not installed")
class TestFormatPatternDetection(unittest.TestCase):
    """Test detection of formatting patterns in documents."""

    def setUp(self):
        """Create test document with various formatting."""
        self.temp_dir = tempfile.mkdtemp()
        self.test_doc_path = Path(self.temp_dir) / "test_formats.docx"

        doc = Document()

        # Add normal paragraph
        doc.add_paragraph("Normal text")

        # Add bold text
        p = doc.add_paragraph()
        run = p.add_run("Bold text")
        run.bold = True

        # Add italic text
        p = doc.add_paragraph()
        run = p.add_run("Italic text")
        run.italic = True

        # Add bullet list
        doc.add_paragraph("Bullet item 1", style='List Bullet')
        doc.add_paragraph("Bullet item 2", style='List Bullet')

        # Add table with text
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Data 1"
        table.cell(1, 1).text = "Data 2"

        doc.save(str(self.test_doc_path))

    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_detect_bold_pattern(self):
        """Should detect bold formatting pattern."""
        from docx_style_creator import detect_formatting_patterns

        patterns = detect_formatting_patterns(self.test_doc_path)

        self.assertIn('bold', patterns)
        self.assertGreater(patterns['bold']['count'], 0)

    def test_detect_italic_pattern(self):
        """Should detect italic formatting pattern."""
        from docx_style_creator import detect_formatting_patterns

        patterns = detect_formatting_patterns(self.test_doc_path)

        self.assertIn('italic', patterns)
        self.assertGreater(patterns['italic']['count'], 0)

    def test_detect_list_pattern(self):
        """Should detect list/bullet formatting pattern."""
        from docx_style_creator import detect_formatting_patterns

        patterns = detect_formatting_patterns(self.test_doc_path)

        self.assertIn('list', patterns)
        self.assertGreater(patterns['list']['count'], 0)

    def test_patterns_include_formatting_details(self):
        """Detected patterns should include formatting details."""
        from docx_style_creator import detect_formatting_patterns

        patterns = detect_formatting_patterns(self.test_doc_path)

        # Bold pattern should have details
        if 'bold' in patterns:
            self.assertIn('count', patterns['bold'])


@unittest.skipUnless(DOCX_AVAILABLE, "python-docx not installed")
class TestStyleCreation(unittest.TestCase):
    """Test creation of styles from detected patterns."""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()
        self.test_doc_path = Path(self.temp_dir) / "test_create.docx"

        doc = Document()
        doc.add_paragraph("Test")
        doc.save(str(self.test_doc_path))

    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_create_paragraph_style(self):
        """Should create a new paragraph style."""
        from docx_style_creator import create_style

        doc = Document(str(self.test_doc_path))

        style = create_style(
            doc,
            name="Test Para Style",
            style_type="paragraph",
            based_on="Normal"
        )

        self.assertIsNotNone(style)
        self.assertEqual(style.name, "Test Para Style")

    def test_create_character_style(self):
        """Should create a new character style."""
        from docx_style_creator import create_style

        doc = Document(str(self.test_doc_path))

        style = create_style(
            doc,
            name="Test Char Style",
            style_type="character",
            bold=True
        )

        self.assertIsNotNone(style)
        self.assertEqual(style.name, "Test Char Style")

    def test_create_style_with_font(self):
        """Should create style with font settings."""
        from docx_style_creator import create_style

        doc = Document(str(self.test_doc_path))

        style = create_style(
            doc,
            name="Custom Font Style",
            style_type="paragraph",
            font_name="Arial",
            font_size=12
        )

        self.assertIsNotNone(style)
        # Verify font was set
        self.assertEqual(style.font.name, "Arial")
        self.assertEqual(style.font.size.pt, 12)

    def test_create_bold_style(self):
        """Should create a bold character style."""
        from docx_style_creator import create_style

        doc = Document(str(self.test_doc_path))

        style = create_style(
            doc,
            name="Strong",
            style_type="character",
            bold=True
        )

        self.assertTrue(style.font.bold)


@unittest.skipUnless(DOCX_AVAILABLE, "python-docx not installed")
class TestStyleApplication(unittest.TestCase):
    """Test applying styles to matching content."""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()
        self.test_doc_path = Path(self.temp_dir) / "test_apply.docx"
        self.output_path = Path(self.temp_dir) / "output.docx"

        doc = Document()
        # Add paragraph with bold run
        p = doc.add_paragraph()
        p.add_run("Normal ")
        bold_run = p.add_run("BOLD")
        bold_run.bold = True
        p.add_run(" normal")

        doc.save(str(self.test_doc_path))

    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_apply_style_to_bold_runs(self):
        """Should apply character style to bold runs."""
        from docx_style_creator import apply_style_to_pattern

        doc = Document(str(self.test_doc_path))

        # Create the Strong style
        from docx_style_creator import create_style
        create_style(doc, name="Strong", style_type="character", bold=True)

        # Apply to bold runs
        count = apply_style_to_pattern(doc, pattern_type="bold", style_name="Strong")

        self.assertGreater(count, 0)


@unittest.skipUnless(DOCX_AVAILABLE, "python-docx not installed")
class TestFullWorkflow(unittest.TestCase):
    """Test complete detect -> create -> apply workflow."""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()
        self.test_doc_path = Path(self.temp_dir) / "test_workflow.docx"
        self.output_path = Path(self.temp_dir) / "output.docx"

        doc = Document()

        # Create document with various formatting
        doc.add_heading("Test Document", level=1)

        p = doc.add_paragraph()
        p.add_run("Some ")
        bold = p.add_run("bold text")
        bold.bold = True
        p.add_run(" here.")

        doc.add_paragraph("Bullet 1", style='List Bullet')
        doc.add_paragraph("Bullet 2", style='List Bullet')

        doc.save(str(self.test_doc_path))

    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_full_workflow(self):
        """Should detect patterns, create styles, and apply them."""
        from docx_style_creator import (
            detect_formatting_patterns,
            create_style,
            apply_style_to_pattern,
            process_document
        )

        # Full workflow function
        result = process_document(
            self.test_doc_path,
            self.output_path,
            create_styles=['Strong']
        )

        self.assertTrue(self.output_path.exists())
        self.assertIn('styles_created', result)
        self.assertIn('patterns_applied', result)


if __name__ == "__main__":
    unittest.main()
