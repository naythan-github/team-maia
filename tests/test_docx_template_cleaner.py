#!/usr/bin/env python3
"""
Tests for Word Document Template Cleaner.

Tests style filtering, latentStyles cleanup, and template generation.
"""

import tempfile
import unittest
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

import sys
sys.path.insert(0, str(Path(__file__).parent.parent / "claude" / "tools" / "document"))

from docx_template_cleaner import (
    get_used_styles_from_document,
    get_style_dependencies,
    get_essential_styles,
    clean_latent_styles,
    filter_styles_xml,
    clean_template,
    NAMESPACES,
)


@unittest.skipUnless(DOCX_AVAILABLE, "python-docx not installed")
class TestUsedStyleExtraction(unittest.TestCase):
    """Test extraction of used styles from document content."""

    def setUp(self):
        """Create test document."""
        self.temp_dir = tempfile.mkdtemp()
        self.test_doc_path = Path(self.temp_dir) / "test_used.docx"

        doc = Document()
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph("Normal text")
        doc.add_paragraph("More text")
        doc.save(str(self.test_doc_path))

    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_returns_categorized_styles(self):
        """Should return dict with paragraph, character, table keys."""
        used = get_used_styles_from_document(self.test_doc_path)

        self.assertIn("paragraph", used)
        self.assertIn("character", used)
        self.assertIn("table", used)

    def test_detects_heading_style(self):
        """Should detect Heading 1 style."""
        used = get_used_styles_from_document(self.test_doc_path)

        # Check for Heading 1 (name or ID)
        para_styles = used["paragraph"]
        has_heading = any("Heading" in s for s in para_styles)
        self.assertTrue(has_heading)

    def test_detects_normal_style(self):
        """Should detect Normal paragraph style."""
        used = get_used_styles_from_document(self.test_doc_path)

        para_styles = used["paragraph"]
        self.assertIn("Normal", para_styles)


class TestEssentialStyles(unittest.TestCase):
    """Test essential style list."""

    def test_essential_styles_includes_normal(self):
        """Essential styles must include Normal."""
        essential = get_essential_styles()
        self.assertIn("Normal", essential)

    def test_essential_styles_includes_default_paragraph_font(self):
        """Essential styles must include DefaultParagraphFont."""
        essential = get_essential_styles()
        self.assertIn("DefaultParagraphFont", essential)

    def test_essential_styles_includes_table_normal(self):
        """Essential styles must include TableNormal."""
        essential = get_essential_styles()
        self.assertIn("TableNormal", essential)


class TestLatentStylesCleaning(unittest.TestCase):
    """Test latentStyles XML cleaning."""

    def test_clean_latent_styles_sets_defaults(self):
        """Should set restrictive defaults."""
        w = '{' + NAMESPACES['w'] + '}'

        # Create mock latentStyles element
        latent = ET.Element(f'{w}latentStyles')

        cleaned = clean_latent_styles(latent, {"Normal", "Heading 1"})

        # Check defaults are restrictive
        self.assertEqual(cleaned.get(f'{w}defSemiHidden'), '1')
        self.assertEqual(cleaned.get(f'{w}defUnhideWhenUsed'), '1')

    def test_clean_latent_styles_creates_exceptions(self):
        """Should create exceptions for keep styles."""
        w = '{' + NAMESPACES['w'] + '}'

        latent = ET.Element(f'{w}latentStyles')
        keep_styles = {"Normal", "Custom Style"}

        cleaned = clean_latent_styles(latent, keep_styles)

        # Check exceptions exist
        exceptions = cleaned.findall(f'{w}lsdException')
        exception_names = {e.get(f'{w}name') for e in exceptions}

        self.assertIn("Normal", exception_names)

    def test_clean_latent_styles_includes_default_visible(self):
        """Should include default visible styles (Heading 1, 2, 3, etc)."""
        w = '{' + NAMESPACES['w'] + '}'

        latent = ET.Element(f'{w}latentStyles')

        cleaned = clean_latent_styles(latent, set())

        exceptions = cleaned.findall(f'{w}lsdException')
        exception_names = {e.get(f'{w}name') for e in exceptions}

        # Should have standard visible styles
        self.assertIn("Heading 1", exception_names)
        self.assertIn("Heading 2", exception_names)
        self.assertIn("Normal", exception_names)


@unittest.skipUnless(DOCX_AVAILABLE, "python-docx not installed")
class TestFullTemplateCleaning(unittest.TestCase):
    """Test complete template cleaning workflow."""

    def setUp(self):
        """Create test document with multiple styles."""
        self.temp_dir = tempfile.mkdtemp()
        self.input_path = Path(self.temp_dir) / "input.docx"
        self.output_path = Path(self.temp_dir) / "output.docx"

        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_heading("Section", level=2)
        doc.add_paragraph("Content paragraph.")
        doc.save(str(self.input_path))

    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_clean_template_creates_output(self):
        """Cleaning should create output file."""
        stats = clean_template(self.input_path, self.output_path)

        self.assertTrue(self.output_path.exists())
        self.assertIn("kept_styles", stats)
        self.assertIn("removed_styles", stats)

    def test_clean_template_reduces_styles(self):
        """Cleaned template should have fewer styles."""
        stats = clean_template(self.input_path, self.output_path)

        self.assertGreater(stats["original_styles"], stats["kept_styles"])

    def test_cleaned_document_opens_in_python_docx(self):
        """Cleaned document should be valid and openable."""
        clean_template(self.input_path, self.output_path)

        # This should not raise
        doc = Document(str(self.output_path))
        self.assertIsNotNone(doc)

    def test_cleaned_document_is_valid_zip(self):
        """Cleaned document should be a valid ZIP/OOXML package."""
        clean_template(self.input_path, self.output_path)

        self.assertTrue(zipfile.is_zipfile(self.output_path))

        with zipfile.ZipFile(self.output_path) as zf:
            self.assertIn("word/document.xml", zf.namelist())
            self.assertIn("word/styles.xml", zf.namelist())

    def test_dry_run_does_not_create_file(self):
        """Dry run should not create output file."""
        stats = clean_template(self.input_path, self.output_path, dry_run=True)

        self.assertFalse(self.output_path.exists())
        self.assertTrue(stats["dry_run"])


@unittest.skipUnless(DOCX_AVAILABLE, "python-docx not installed")
class TestStyleDependencies(unittest.TestCase):
    """Test style dependency resolution."""

    def setUp(self):
        """Create test document."""
        self.temp_dir = tempfile.mkdtemp()
        self.test_doc_path = Path(self.temp_dir) / "test_deps.docx"

        doc = Document()
        doc.add_heading("Test", level=2)  # Heading 2 depends on Heading 1
        doc.save(str(self.test_doc_path))

    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_dependencies_include_base_styles(self):
        """Should include base styles in dependency chain."""
        with zipfile.ZipFile(self.test_doc_path) as zf:
            styles_xml = zf.read("word/styles.xml")

        styles_root = ET.fromstring(styles_xml)

        # Request Heading2, should get its base style too
        deps = get_style_dependencies(styles_root, {"Heading2"})

        # Should have at least the requested style
        self.assertIn("Heading2", deps)


if __name__ == "__main__":
    unittest.main()
