#!/usr/bin/env python3
"""
Tests for Word Document Style Analyzer.

Tests style extraction, usage counting, and analysis reporting.
"""

import json
import tempfile
import unittest
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

import sys
sys.path.insert(0, str(Path(__file__).parent.parent / "claude" / "tools" / "document"))

from docx_style_analyzer import (
    get_all_defined_styles,
    analyze_paragraph_styles,
    analyze_character_styles,
    analyze_table_styles,
    analyze_document,
)


@unittest.skipUnless(DOCX_AVAILABLE, "python-docx not installed")
class TestStyleExtraction(unittest.TestCase):
    """Test style definition extraction."""

    def setUp(self):
        """Create a test document with known styles."""
        self.temp_dir = tempfile.mkdtemp()
        self.test_doc_path = Path(self.temp_dir) / "test_styles.docx"

        doc = Document()
        # Add content with different styles
        doc.add_heading("Test Heading 1", level=1)
        doc.add_heading("Test Heading 2", level=2)
        doc.add_paragraph("Normal paragraph text.")
        doc.add_paragraph("Another normal paragraph.")
        doc.save(str(self.test_doc_path))

    def tearDown(self):
        """Clean up test files."""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_get_defined_styles_returns_dict(self):
        """Defined styles should return a dictionary."""
        doc = Document(str(self.test_doc_path))
        styles = get_all_defined_styles(doc)
        self.assertIsInstance(styles, dict)
        self.assertGreater(len(styles), 0)

    def test_defined_styles_contain_normal(self):
        """Normal style should always be defined."""
        doc = Document(str(self.test_doc_path))
        styles = get_all_defined_styles(doc)
        self.assertIn("Normal", styles)

    def test_style_info_has_required_fields(self):
        """Each style should have name, type, builtin fields."""
        doc = Document(str(self.test_doc_path))
        styles = get_all_defined_styles(doc)
        for name, info in styles.items():
            self.assertIn("name", info)
            self.assertIn("type", info)
            self.assertIn("builtin", info)


@unittest.skipUnless(DOCX_AVAILABLE, "python-docx not installed")
class TestParagraphStyleAnalysis(unittest.TestCase):
    """Test paragraph style usage analysis."""

    def setUp(self):
        """Create test document."""
        self.temp_dir = tempfile.mkdtemp()
        self.test_doc_path = Path(self.temp_dir) / "test_para.docx"

        doc = Document()
        doc.add_heading("Heading", level=1)
        doc.add_paragraph("Para 1")
        doc.add_paragraph("Para 2")
        doc.add_paragraph("Para 3")
        doc.save(str(self.test_doc_path))

    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_paragraph_styles_counted(self):
        """Paragraph styles should be counted correctly."""
        doc = Document(str(self.test_doc_path))
        usage = analyze_paragraph_styles(doc)
        self.assertIsInstance(usage, dict)
        # Should have Normal and Heading 1
        self.assertIn("Normal", usage)
        self.assertEqual(usage["Normal"], 3)  # 3 normal paragraphs

    def test_heading_style_detected(self):
        """Heading styles should be detected."""
        doc = Document(str(self.test_doc_path))
        usage = analyze_paragraph_styles(doc)
        self.assertIn("Heading 1", usage)
        self.assertEqual(usage["Heading 1"], 1)


@unittest.skipUnless(DOCX_AVAILABLE, "python-docx not installed")
class TestTableStyleAnalysis(unittest.TestCase):
    """Test table style usage analysis."""

    def setUp(self):
        """Create test document with table."""
        self.temp_dir = tempfile.mkdtemp()
        self.test_doc_path = Path(self.temp_dir) / "test_table.docx"

        doc = Document()
        doc.add_paragraph("Before table")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        doc.save(str(self.test_doc_path))

    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_table_styles_counted(self):
        """Table styles should be counted."""
        doc = Document(str(self.test_doc_path))
        usage = analyze_table_styles(doc)
        self.assertIsInstance(usage, dict)
        self.assertEqual(len(usage), 1)  # One table


@unittest.skipUnless(DOCX_AVAILABLE, "python-docx not installed")
class TestFullAnalysis(unittest.TestCase):
    """Test complete document analysis."""

    def setUp(self):
        """Create comprehensive test document."""
        self.temp_dir = tempfile.mkdtemp()
        self.test_doc_path = Path(self.temp_dir) / "test_full.docx"

        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_heading("Section", level=2)
        doc.add_paragraph("Content")
        table = doc.add_table(rows=1, cols=2)
        doc.save(str(self.test_doc_path))

    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_analyze_document_returns_complete_report(self):
        """Full analysis should return all required sections."""
        analysis = analyze_document(self.test_doc_path)

        self.assertIn("stats", analysis)
        self.assertIn("defined_styles", analysis)
        self.assertIn("used_styles", analysis)
        self.assertIn("unused_styles", analysis)
        self.assertIn("recommendations", analysis)

    def test_stats_has_counts(self):
        """Stats should include style counts."""
        analysis = analyze_document(self.test_doc_path)
        stats = analysis["stats"]

        self.assertIn("total_defined", stats)
        self.assertIn("total_used", stats)
        self.assertIn("total_unused", stats)
        self.assertGreater(stats["total_defined"], 0)

    def test_used_styles_list(self):
        """Used styles should be a list with required fields."""
        analysis = analyze_document(self.test_doc_path)
        used = analysis["used_styles"]

        self.assertIsInstance(used, list)
        self.assertGreater(len(used), 0)
        for style in used:
            self.assertIn("name", style)
            self.assertIn("count", style)
            self.assertIn("type", style)


if __name__ == "__main__":
    unittest.main()
