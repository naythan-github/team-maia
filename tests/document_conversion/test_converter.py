#!/usr/bin/env python3
"""
Automated Test Suite for Document Conversion System

Tests Orro corporate styling, color preservation, font application,
and overall conversion quality.

Agent: SRE Principal Engineer Agent
Phase: 163 - Production Hardening
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from orro_md_to_docx import convert_md_to_docx

# Test constants
ORRO_PURPLE = RGBColor(112, 48, 160)
TEST_OUTPUT_DIR = Path("/tmp/maia_converter_tests")


def setup_tests():
    """Create test directory and sample markdown file"""
    TEST_OUTPUT_DIR.mkdir(exist_ok=True)

    # Create test markdown
    test_md = TEST_OUTPUT_DIR / "test_document.md"
    test_md.write_text("""# Heading 1 Test

## Heading 2 Test

### Heading 3 Test

This is a normal paragraph with **bold** and *italic* text.

## Another Heading 2

- Bullet point 1
- Bullet point 2
- Bullet point 3

### Another Heading 3

1. Numbered item 1
2. Numbered item 2
3. Numbered item 3

## Table Test

| Column 1 | Column 2 | Column 3 |
|----------|----------|----------|
| Data 1   | Data 2   | Data 3   |
| Data 4   | Data 5   | Data 6   |

## Conclusion

This is the final paragraph.
""")

    return test_md


def test_color_preservation():
    """Test that Orro purple color is applied to all headings"""
    print("\n🧪 TEST 1: Color Preservation")
    print("-" * 50)

    test_md = setup_tests()
    output_docx = TEST_OUTPUT_DIR / "test_color.docx"

    # Convert
    success = convert_md_to_docx(test_md, output_docx, apply_table_styles=True)

    if not success:
        print("   ❌ FAILED: Conversion failed")
        return False

    # Validate colors
    doc = Document(output_docx)
    headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]

    if not headings:
        print("   ❌ FAILED: No headings found")
        return False

    purple_count = 0
    total_headings = len(headings)

    for heading in headings:
        if heading.runs:
            for run in heading.runs:
                if (hasattr(run.font, 'color') and run.font.color and
                    hasattr(run.font.color, 'rgb') and run.font.color.rgb):
                    rgb = run.font.color.rgb
                    if rgb == ORRO_PURPLE:
                        purple_count += 1
                        break

    if purple_count == total_headings:
        print(f"   ✅ PASSED: {purple_count}/{total_headings} headings have Orro purple")
        return True
    else:
        print(f"   ❌ FAILED: Only {purple_count}/{total_headings} headings have Orro purple")
        return False


def test_structure_preservation():
    """Test that document structure is preserved (headings, paragraphs, lists)"""
    print("\n🧪 TEST 2: Structure Preservation")
    print("-" * 50)

    test_md = setup_tests()
    output_docx = TEST_OUTPUT_DIR / "test_structure.docx"

    # Convert
    success = convert_md_to_docx(test_md, output_docx, apply_table_styles=True)

    if not success:
        print("   ❌ FAILED: Conversion failed")
        return False

    # Validate structure
    doc = Document(output_docx)

    h1_count = len([p for p in doc.paragraphs if p.style.name == 'Heading 1'])
    h2_count = len([p for p in doc.paragraphs if p.style.name == 'Heading 2'])
    h3_count = len([p for p in doc.paragraphs if p.style.name == 'Heading 3'])
    table_count = len(doc.tables)

    # Expected structure from test markdown
    expected = {
        'h1': 1,
        'h2': 4,
        'h3': 2,
        'tables': 1
    }

    passed = True

    if h1_count == expected['h1']:
        print(f"   ✅ H1 count: {h1_count} (expected {expected['h1']})")
    else:
        print(f"   ❌ H1 count: {h1_count} (expected {expected['h1']})")
        passed = False

    if h2_count == expected['h2']:
        print(f"   ✅ H2 count: {h2_count} (expected {expected['h2']})")
    else:
        print(f"   ❌ H2 count: {h2_count} (expected {expected['h2']})")
        passed = False

    if h3_count == expected['h3']:
        print(f"   ✅ H3 count: {h3_count} (expected {expected['h3']})")
    else:
        print(f"   ❌ H3 count: {h3_count} (expected {expected['h3']})")
        passed = False

    if table_count == expected['tables']:
        print(f"   ✅ Table count: {table_count} (expected {expected['tables']})")
    else:
        print(f"   ❌ Table count: {table_count} (expected {expected['tables']})")
        passed = False

    return passed


def test_margin_formatting():
    """Test that 1.0 inch margins are applied"""
    print("\n🧪 TEST 3: Margin Formatting")
    print("-" * 50)

    test_md = setup_tests()
    output_docx = TEST_OUTPUT_DIR / "test_margins.docx"

    # Convert
    success = convert_md_to_docx(test_md, output_docx, apply_table_styles=True)

    if not success:
        print("   ❌ FAILED: Conversion failed")
        return False

    # Validate margins
    doc = Document(output_docx)
    section = doc.sections[0]

    expected_margin_inches = 1.0
    tolerance = 0.01  # Allow small floating point differences

    margins = {
        'top': section.top_margin.inches,
        'bottom': section.bottom_margin.inches,
        'left': section.left_margin.inches,
        'right': section.right_margin.inches
    }

    passed = True

    for name, value in margins.items():
        if abs(value - expected_margin_inches) < tolerance:
            print(f"   ✅ {name.capitalize()} margin: {value:.2f}\" (expected {expected_margin_inches}\")")
        else:
            print(f"   ❌ {name.capitalize()} margin: {value:.2f}\" (expected {expected_margin_inches}\")")
            passed = False

    return passed


def test_performance():
    """Test that conversion completes within acceptable time (<5s)"""
    print("\n🧪 TEST 4: Performance")
    print("-" * 50)

    import time

    test_md = setup_tests()
    output_docx = TEST_OUTPUT_DIR / "test_performance.docx"

    # Measure conversion time
    start_time = time.time()
    success = convert_md_to_docx(test_md, output_docx, apply_table_styles=True)
    elapsed_time = time.time() - start_time

    if not success:
        print("   ❌ FAILED: Conversion failed")
        return False

    target_time = 5.0  # seconds

    if elapsed_time < target_time:
        print(f"   ✅ PASSED: Conversion took {elapsed_time:.2f}s (target: <{target_time}s)")
        return True
    else:
        print(f"   ❌ FAILED: Conversion took {elapsed_time:.2f}s (target: <{target_time}s)")
        return False


def run_all_tests():
    """Run all tests and report results"""
    print("=" * 70)
    print("DOCUMENT CONVERTER - AUTOMATED TEST SUITE")
    print("=" * 70)

    tests = [
        ("Color Preservation", test_color_preservation),
        ("Structure Preservation", test_structure_preservation),
        ("Margin Formatting", test_margin_formatting),
        ("Performance", test_performance)
    ]

    results = []

    for test_name, test_func in tests:
        try:
            passed = test_func()
            results.append((test_name, passed))
        except Exception as e:
            print(f"\n   ❌ EXCEPTION: {e}")
            results.append((test_name, False))

    # Summary
    print("\n" + "=" * 70)
    print("TEST SUMMARY")
    print("=" * 70)

    passed_count = sum(1 for _, passed in results if passed)
    total_count = len(results)

    for test_name, passed in results:
        status = "✅ PASSED" if passed else "❌ FAILED"
        print(f"   {status}: {test_name}")

    print()
    print(f"Results: {passed_count}/{total_count} tests passed ({passed_count/total_count*100:.1f}%)")
    print("=" * 70)

    # Cleanup
    print(f"\n📁 Test files saved in: {TEST_OUTPUT_DIR}")

    return passed_count == total_count


if __name__ == "__main__":
    success = run_all_tests()
    sys.exit(0 if success else 1)
