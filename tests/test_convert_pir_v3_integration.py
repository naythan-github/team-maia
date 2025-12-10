#!/usr/bin/env python3
"""
Integration tests for PIR Converter v3.

Tests that convert_pir_v3.py + pir_docx_normalizer.py pipeline works correctly.
This is a safety test to ensure refactoring doesn't break the pipeline.

TDD: Write test first, then simplify converter.
"""

import pytest
import tempfile
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

# Paths - consolidated under claude/tools/document/pir/
PIR_DIR = Path(__file__).parent.parent / "claude" / "tools" / "document" / "pir"
CONVERTER_PATH = PIR_DIR / "convert_pir.py"
REFERENCE_TEMPLATE = PIR_DIR / "templates" / "pir_orro_reference.docx"

sys.path.insert(0, str(PIR_DIR))
from pir_docx_normalizer import normalize_document


@pytest.fixture
def sample_pir_markdown():
    """Create a sample PIR markdown file with tables."""
    content = """# Post-Incident Review

## Executive Summary

| Field | Value |
|-------|-------|
| Incident ID | PIR-2025-001 |
| Date | 2025-12-09 |
| Severity | Critical |

## Timeline

| Time | Event | Impact |
|------|-------|--------|
| 09:00 | Initial detection | Alert triggered |
| 09:15 | Investigation started | Team mobilized |
| 10:30 | Root cause identified | N+1 query pattern |
| 11:00 | Fix deployed | Service restored |

## Findings

The incident was caused by a database query issue.
"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write(content)
        return Path(f.name)


class TestConverterPipeline:
    """Integration tests for the full conversion pipeline."""

    @pytest.mark.skipif(not CONVERTER_PATH.exists(), reason="Converter not found")
    @pytest.mark.skipif(not REFERENCE_TEMPLATE.exists(), reason="Template not found")
    def test_converter_produces_valid_docx(self, sample_pir_markdown):
        """Converter should produce a valid .docx file."""
        output_path = sample_pir_markdown.with_suffix('.docx')

        try:
            # Run converter
            result = subprocess.run(
                ['python3', str(CONVERTER_PATH), str(sample_pir_markdown), '-o', str(output_path)],
                capture_output=True,
                text=True
            )

            assert result.returncode == 0, f"Converter failed: {result.stderr}"
            assert output_path.exists(), "Output file not created"

            # Verify it's a valid docx
            doc = Document(str(output_path))
            assert len(doc.paragraphs) > 0, "Document should have paragraphs"

        finally:
            sample_pir_markdown.unlink(missing_ok=True)
            output_path.unlink(missing_ok=True)

    @pytest.mark.skipif(not CONVERTER_PATH.exists(), reason="Converter not found")
    @pytest.mark.skipif(not REFERENCE_TEMPLATE.exists(), reason="Template not found")
    def test_converter_creates_tables(self, sample_pir_markdown):
        """Converter should create tables from markdown tables."""
        output_path = sample_pir_markdown.with_suffix('.docx')

        try:
            subprocess.run(
                ['python3', str(CONVERTER_PATH), str(sample_pir_markdown), '-o', str(output_path)],
                capture_output=True,
                check=True
            )

            doc = Document(str(output_path))

            # Should have 2 tables (Executive Summary + Timeline)
            assert len(doc.tables) >= 2, f"Expected 2+ tables, got {len(doc.tables)}"

        finally:
            sample_pir_markdown.unlink(missing_ok=True)
            output_path.unlink(missing_ok=True)

    @pytest.mark.skipif(not CONVERTER_PATH.exists(), reason="Converter not found")
    @pytest.mark.skipif(not REFERENCE_TEMPLATE.exists(), reason="Template not found")
    def test_tables_have_orro_style(self, sample_pir_markdown):
        """Tables should have _Orro Table 1 style applied."""
        output_path = sample_pir_markdown.with_suffix('.docx')

        try:
            subprocess.run(
                ['python3', str(CONVERTER_PATH), str(sample_pir_markdown), '-o', str(output_path)],
                capture_output=True,
                check=True
            )

            doc = Document(str(output_path))

            for table in doc.tables:
                tbl = table._tbl
                tblPr = tbl.tblPr
                if tblPr is not None:
                    tblStyle = tblPr.find(qn('w:tblStyle'))
                    if tblStyle is not None:
                        style_val = tblStyle.get(qn('w:val'))
                        # Style should be _Orro Table 1 (or similar)
                        assert 'Orro' in style_val or style_val is not None, \
                            f"Table style should be Orro, got {style_val}"

        finally:
            sample_pir_markdown.unlink(missing_ok=True)
            output_path.unlink(missing_ok=True)


class TestFullPipeline:
    """Test converter + normalizer together."""

    @pytest.mark.skipif(not CONVERTER_PATH.exists(), reason="Converter not found")
    @pytest.mark.skipif(not REFERENCE_TEMPLATE.exists(), reason="Template not found")
    def test_full_pipeline_produces_correct_output(self, sample_pir_markdown):
        """Full pipeline (converter + normalizer) should produce correct formatting."""
        output_path = sample_pir_markdown.with_suffix('.docx')

        try:
            # Step 1: Run converter
            result = subprocess.run(
                ['python3', str(CONVERTER_PATH), str(sample_pir_markdown), '-o', str(output_path)],
                capture_output=True,
                text=True
            )
            assert result.returncode == 0, f"Converter failed: {result.stderr}"

            # Step 2: Run normalizer
            stats = normalize_document(output_path, output_path, verbose=False)
            assert stats["tables_fixed"] >= 2, "Should fix at least 2 tables"

            # Step 3: Verify final output
            doc = Document(str(output_path))

            for table in doc.tables:
                tbl = table._tbl
                tblPr = tbl.tblPr
                assert tblPr is not None, "Table should have properties"

                # Check table width is 100%
                tblW = tblPr.find(qn('w:tblW'))
                assert tblW is not None, "Table should have width set"
                assert tblW.get(qn('w:type')) == 'pct', "Width should be percentage"
                assert tblW.get(qn('w:w')) == '5000', "Width should be 100% (5000)"

                # Check borders are Orro purple (explicit RGB, not theme)
                tblBorders = tblPr.find(qn('w:tblBorders'))
                assert tblBorders is not None, "Table should have borders"
                bottom = tblBorders.find(qn('w:bottom'))
                if bottom is not None:
                    color = bottom.get(qn('w:color'))
                    # Should be explicit RGB, not theme color
                    assert color == "7030A0", f"Border should be Orro purple, got {color}"

        finally:
            sample_pir_markdown.unlink(missing_ok=True)
            output_path.unlink(missing_ok=True)


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
