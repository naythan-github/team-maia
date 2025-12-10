#!/usr/bin/env python3
"""
TDD tests for PIR Document Normalizer.

Tests validate that:
1. Tables are set to 100% width
2. Original cell width RATIOS are preserved (not forced to equal)
3. Cell margins are applied consistently
4. Document structure is not corrupted
"""

import pytest
import tempfile
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Inches, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
sys.path.insert(0, str(Path(__file__).parent.parent / "claude" / "tools" / "document"))
from pir_docx_normalizer import normalize_document, calculate_content_aware_widths


# --- Test Fixtures ---

@pytest.fixture
def sample_doc_with_tables():
    """Create a test document with tables of varying column widths."""
    doc = Document()

    # Table 1: 2 columns with UNEQUAL widths (30% / 70% ratio)
    # This simulates a label + content layout
    table1 = doc.add_table(rows=2, cols=2)
    table1.style = 'Table Grid'
    table1.rows[0].cells[0].text = "Label"
    table1.rows[0].cells[1].text = "Content value that should be wider"

    # Set explicit unequal widths (in twips: 1440 = 1 inch)
    # Col 1: 2 inches, Col 2: 5 inches (2:5 ratio)
    for row in table1.rows:
        set_cell_width_dxa(row.cells[0], 2880)  # 2 inches
        set_cell_width_dxa(row.cells[1], 7200)  # 5 inches

    doc.add_paragraph("")

    # Table 2: 3 equal columns
    table2 = doc.add_table(rows=2, cols=3)
    table2.style = 'Table Grid'
    for i, cell in enumerate(table2.rows[0].cells):
        cell.text = f"Col {i+1}"
        set_cell_width_dxa(cell, 2880)  # All equal

    doc.add_paragraph("")

    # Table 3: 4 columns with varied widths (1:2:2:1 ratio)
    table3 = doc.add_table(rows=2, cols=4)
    table3.style = 'Table Grid'
    widths = [1440, 2880, 2880, 1440]  # 1:2:2:1 ratio
    for row in table3.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width_dxa(cell, widths[i])

    # Save to temp file
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        doc.save(f.name)
        return Path(f.name)


def set_cell_width_dxa(cell, width_twips):
    """Helper to set cell width in twips (dxa)."""
    tc = cell._tc
    tcPr = tc.tcPr
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)

    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.insert(0, tcW)

    tcW.set(qn('w:type'), 'dxa')
    tcW.set(qn('w:w'), str(width_twips))


def get_table_width(table):
    """Get table width type and value."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is not None:
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is not None:
            return {
                "type": tblW.get(qn('w:type')),
                "value": tblW.get(qn('w:w'))
            }
    return {"type": None, "value": None}


def get_cell_widths(table):
    """Get cell width ratios for first row."""
    widths = []
    if table.rows:
        for cell in table.rows[0].cells:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is not None:
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    widths.append({
                        "type": tcW.get(qn('w:type')),
                        "value": int(tcW.get(qn('w:w')) or 0)
                    })
    return widths


def calculate_ratios(widths):
    """Calculate width ratios from cell widths."""
    values = [w["value"] for w in widths]
    total = sum(values)
    if total == 0:
        return []
    return [round(v / total, 2) for v in values]


# --- Tests ---

class TestTableWidth:
    """Tests for table width normalization."""

    def test_tables_set_to_100_percent(self, sample_doc_with_tables):
        """All tables should be set to 100% width after normalization."""
        output_path = sample_doc_with_tables.with_name("output.docx")

        normalize_document(sample_doc_with_tables, output_path, verbose=False)

        doc = Document(str(output_path))
        for i, table in enumerate(doc.tables):
            width = get_table_width(table)
            assert width["type"] == "pct", f"Table {i}: width type should be 'pct', got '{width['type']}'"
            assert width["value"] == "5000", f"Table {i}: width should be 5000 (100%), got '{width['value']}'"

        # Cleanup
        output_path.unlink()
        sample_doc_with_tables.unlink()


class TestCellWidthPreservation:
    """Tests for preserving cell width ratios."""

    def test_unequal_column_ratios_preserved(self, sample_doc_with_tables):
        """
        CRITICAL: Unequal column widths must maintain their RATIO.
        A 2:5 ratio table should NOT become 1:1.
        """
        # Get original ratios BEFORE normalization
        doc_before = Document(str(sample_doc_with_tables))
        original_ratios = calculate_ratios(get_cell_widths(doc_before.tables[0]))

        # Normalize with preserve-ratios mode
        output_path = sample_doc_with_tables.with_name("output.docx")
        normalize_document(sample_doc_with_tables, output_path, verbose=False, content_aware=False)

        # Get ratios AFTER normalization
        doc_after = Document(str(output_path))
        new_ratios = calculate_ratios(get_cell_widths(doc_after.tables[0]))

        # Ratios should be approximately the same (within 5% tolerance)
        assert len(original_ratios) == len(new_ratios), "Column count changed"
        for i, (orig, new) in enumerate(zip(original_ratios, new_ratios)):
            assert abs(orig - new) < 0.05, \
                f"Column {i} ratio changed: {orig} -> {new}. Original ratios destroyed!"

        # Cleanup
        output_path.unlink()
        sample_doc_with_tables.unlink()

    def test_equal_columns_stay_equal(self, sample_doc_with_tables):
        """Equal width columns should remain equal after normalization."""
        doc_before = Document(str(sample_doc_with_tables))
        original_ratios = calculate_ratios(get_cell_widths(doc_before.tables[1]))

        output_path = sample_doc_with_tables.with_name("output.docx")
        normalize_document(sample_doc_with_tables, output_path, verbose=False, content_aware=False)

        doc_after = Document(str(output_path))
        new_ratios = calculate_ratios(get_cell_widths(doc_after.tables[1]))

        # All should be ~0.33 (equal thirds)
        for ratio in new_ratios:
            assert abs(ratio - 0.33) < 0.05, f"Equal columns became unequal: {new_ratios}"

        # Cleanup
        output_path.unlink()
        sample_doc_with_tables.unlink()

    def test_complex_ratio_preserved(self, sample_doc_with_tables):
        """Complex ratios (1:2:2:1) should be preserved."""
        doc_before = Document(str(sample_doc_with_tables))
        original_ratios = calculate_ratios(get_cell_widths(doc_before.tables[2]))
        # Expected: [0.17, 0.33, 0.33, 0.17] approximately

        output_path = sample_doc_with_tables.with_name("output.docx")
        normalize_document(sample_doc_with_tables, output_path, verbose=False, content_aware=False)

        doc_after = Document(str(output_path))
        new_ratios = calculate_ratios(get_cell_widths(doc_after.tables[2]))

        for i, (orig, new) in enumerate(zip(original_ratios, new_ratios)):
            assert abs(orig - new) < 0.05, \
                f"Column {i} ratio changed: {orig} -> {new}"

        # Cleanup
        output_path.unlink()
        sample_doc_with_tables.unlink()


class TestRealTemplate:
    """Tests against the actual PIR template."""

    def test_template_tables_not_corrupted(self):
        """PIR template cell ratios should be preserved."""
        template_path = Path("/Users/naythandawe/git/maia/claude/tools/security/pir_templates/pir_credential_stuffing_template.docx")

        if not template_path.exists():
            pytest.skip("Template not found")

        # Get original ratios for all tables
        doc_before = Document(str(template_path))
        original_data = []
        for table in doc_before.tables:
            original_data.append({
                "ratios": calculate_ratios(get_cell_widths(table)),
                "rows": len(table.rows),
                "cols": len(table.columns)
            })

        # Normalize to temp file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = Path(f.name)

        shutil.copy(template_path, output_path)
        normalize_document(output_path, output_path, verbose=False, content_aware=False)

        # Verify ratios preserved
        doc_after = Document(str(output_path))

        assert len(doc_after.tables) == len(original_data), "Table count changed"

        for i, (table, orig) in enumerate(zip(doc_after.tables, original_data)):
            new_ratios = calculate_ratios(get_cell_widths(table))

            # Table should be 100% width
            width = get_table_width(table)
            assert width["type"] == "pct" and width["value"] == "5000", \
                f"Table {i} not 100% width"

            # Ratios should be preserved
            if orig["ratios"]:  # Skip if no width data
                for j, (o, n) in enumerate(zip(orig["ratios"], new_ratios)):
                    assert abs(o - n) < 0.05, \
                        f"Table {i} col {j}: ratio {o} -> {n}"

        # Cleanup
        output_path.unlink()


class TestOrroBorders:
    """Tests for Orro corporate border styling."""

    ORRO_PURPLE = "7030A0"  # Main border color
    ORRO_LIGHT_PURPLE = "CBC6F3"  # Inside horizontal lines

    @pytest.fixture
    def styled_table_doc(self):
        """Create a doc with a table that has Orro style applied."""
        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "Header 1"
        table.rows[0].cells[1].text = "Header 2"
        table.rows[1].cells[0].text = "Data 1"
        table.rows[1].cells[1].text = "Data 2"

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            return Path(f.name)

    def test_orro_borders_applied(self, styled_table_doc):
        """Tables should have Orro purple borders after normalization."""
        from pir_docx_normalizer import normalize_document, apply_orro_borders

        output_path = styled_table_doc.with_name("output.docx")
        normalize_document(styled_table_doc, output_path, verbose=False)

        doc = Document(str(output_path))
        table = doc.tables[0]

        # Check table has borders
        tbl = table._tbl
        tblPr = tbl.tblPr
        assert tblPr is not None, "Table should have properties"

        tblBorders = tblPr.find(qn('w:tblBorders'))
        assert tblBorders is not None, "Table should have borders defined"

        # Check bottom border is Orro purple
        bottom = tblBorders.find(qn('w:bottom'))
        assert bottom is not None, "Table should have bottom border"
        assert bottom.get(qn('w:color')) == self.ORRO_PURPLE, \
            f"Bottom border should be Orro purple, got {bottom.get(qn('w:color'))}"

        # Check inside horizontal is light purple
        insideH = tblBorders.find(qn('w:insideH'))
        assert insideH is not None, "Table should have inside horizontal border"
        assert insideH.get(qn('w:color')) == self.ORRO_LIGHT_PURPLE, \
            f"Inside H should be light purple, got {insideH.get(qn('w:color'))}"

        # Cleanup
        output_path.unlink()
        styled_table_doc.unlink()

    def test_no_vertical_borders(self, styled_table_doc):
        """Orro style should have no vertical borders (clean look)."""
        from pir_docx_normalizer import normalize_document

        output_path = styled_table_doc.with_name("output.docx")
        normalize_document(styled_table_doc, output_path, verbose=False)

        doc = Document(str(output_path))
        table = doc.tables[0]
        tblBorders = table._tbl.tblPr.find(qn('w:tblBorders'))

        # Left, right, insideV should be nil or not present
        for border_name in ['left', 'right', 'insideV']:
            border = tblBorders.find(qn(f'w:{border_name}'))
            if border is not None:
                assert border.get(qn('w:val')) == 'nil', \
                    f"{border_name} border should be nil"

        # Cleanup
        output_path.unlink()
        styled_table_doc.unlink()


class TestContentAwareSizing:
    """Tests for content-aware column width calculation."""

    @pytest.fixture
    def key_value_table_doc(self):
        """Create a doc with a key-value style table (short labels, long values)."""
        doc = Document()
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'

        # Short labels, long values
        data = [
            ("Status", "Active - Production Environment Running Normally"),
            ("ID", "ACC-2025-001234-PROD-CRITICAL-SECURITY-INCIDENT"),
            ("Type", "Security Breach Investigation"),
            ("Date", "2025-12-09"),
        ]
        for i, (label, value) in enumerate(data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            return Path(f.name)

    @pytest.fixture
    def multi_column_data_table_doc(self):
        """Create a doc with a data table with varying content lengths."""
        doc = Document()
        table = doc.add_table(rows=3, cols=5)
        table.style = 'Table Grid'

        # Headers with varying lengths
        headers = ["ID", "Account", "Country", "Status", "Notes"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        # Data rows
        table.rows[1].cells[0].text = "1"
        table.rows[1].cells[1].text = "admin@example.com"
        table.rows[1].cells[2].text = "AU"
        table.rows[1].cells[3].text = "CRITICAL"
        table.rows[1].cells[4].text = "Multiple failed login attempts detected from foreign IP"

        table.rows[2].cells[0].text = "2"
        table.rows[2].cells[1].text = "user@example.com"
        table.rows[2].cells[2].text = "US"
        table.rows[2].cells[3].text = "OK"
        table.rows[2].cells[4].text = "Normal"

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            return Path(f.name)

    def test_key_value_table_proportions(self, key_value_table_doc):
        """Key-value tables should have narrow label column, wide value column."""
        doc = Document(str(key_value_table_doc))
        table = doc.tables[0]

        widths = calculate_content_aware_widths(table)

        # Label column should be significantly narrower than value column
        assert len(widths) == 2
        label_ratio = widths[0] / sum(widths)
        value_ratio = widths[1] / sum(widths)

        # Label should be ~20-35%, value ~65-80%
        assert label_ratio < 0.40, f"Label column too wide: {label_ratio:.0%}"
        assert value_ratio > 0.55, f"Value column too narrow: {value_ratio:.0%}"

        # Cleanup
        key_value_table_doc.unlink()

    def test_multi_column_notes_column_wider(self, multi_column_data_table_doc):
        """Notes/description columns with more content should be wider."""
        doc = Document(str(multi_column_data_table_doc))
        table = doc.tables[0]

        widths = calculate_content_aware_widths(table)

        # 5 columns: ID, Account, Country, Status, Notes
        assert len(widths) == 5

        ratios = [w / sum(widths) for w in widths]

        # ID column (short numbers) should be narrow
        assert ratios[0] < 0.15, f"ID column too wide: {ratios[0]:.0%}"

        # Notes column (longest content) should be widest
        assert ratios[4] > ratios[0], "Notes should be wider than ID"
        assert ratios[4] > ratios[2], "Notes should be wider than Country"

        # Cleanup
        multi_column_data_table_doc.unlink()

    def test_minimum_column_width_enforced(self, key_value_table_doc):
        """No column should be less than minimum width (10%)."""
        doc = Document(str(key_value_table_doc))
        table = doc.tables[0]

        widths = calculate_content_aware_widths(table)
        total = sum(widths)

        for i, w in enumerate(widths):
            ratio = w / total
            assert ratio >= 0.10, f"Column {i} below minimum: {ratio:.0%}"

        # Cleanup
        key_value_table_doc.unlink()

    def test_maximum_column_width_enforced(self, multi_column_data_table_doc):
        """No single column should exceed maximum width (60%)."""
        doc = Document(str(multi_column_data_table_doc))
        table = doc.tables[0]

        widths = calculate_content_aware_widths(table)
        total = sum(widths)

        for i, w in enumerate(widths):
            ratio = w / total
            assert ratio <= 0.60, f"Column {i} exceeds maximum: {ratio:.0%}"

        # Cleanup
        multi_column_data_table_doc.unlink()

    def test_widths_sum_to_100_percent(self, key_value_table_doc):
        """All column widths must sum to exactly 5000 (100%)."""
        doc = Document(str(key_value_table_doc))
        table = doc.tables[0]

        widths = calculate_content_aware_widths(table)

        assert sum(widths) == 5000, f"Widths sum to {sum(widths)}, expected 5000"

        # Cleanup
        key_value_table_doc.unlink()


class TestParagraphSpacing:
    """Tests for paragraph spacing normalization."""

    # Spacing specs (in points)
    BODY_SPACE_AFTER = 6
    HEADING_SPACE_BEFORE = 12
    HEADING_SPACE_AFTER = 6
    LIST_SPACE = 2
    TABLE_CELL_SPACE = 2

    @pytest.fixture
    def doc_with_paragraphs(self):
        """Create a doc with various paragraph styles."""
        from docx.shared import Pt
        doc = Document()

        # Add heading
        doc.add_heading('Test Heading 1', level=1)
        doc.add_heading('Test Heading 2', level=2)

        # Add body paragraphs
        doc.add_paragraph('This is body text paragraph one.')
        doc.add_paragraph('This is body text paragraph two.')

        # Add list items
        doc.add_paragraph('List item one', style='List Bullet')
        doc.add_paragraph('List item two', style='List Bullet')

        # Add table with text
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Header 1"
        table.rows[0].cells[1].text = "Header 2"
        table.rows[1].cells[0].text = "Data 1"
        table.rows[1].cells[1].text = "Data 2"

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            return Path(f.name)

    def test_body_text_spacing(self, doc_with_paragraphs):
        """Body text should have 6pt after spacing."""
        from pir_docx_normalizer import normalize_document
        from docx.shared import Pt

        output_path = doc_with_paragraphs.with_name("output.docx")
        normalize_document(doc_with_paragraphs, output_path, verbose=False)

        doc = Document(str(output_path))

        # Find Normal paragraphs with text
        normal_paras = [p for p in doc.paragraphs
                       if p.style and p.style.name == 'Normal' and p.text.strip()]
        assert len(normal_paras) > 0, "Should have Normal paragraphs to test"

        for para in normal_paras:
            space_after = para.paragraph_format.space_after
            assert space_after is not None, "Body text space_after should be set"
            assert space_after.pt == self.BODY_SPACE_AFTER, \
                f"Body text space_after should be {self.BODY_SPACE_AFTER}pt, got {space_after.pt}pt"

        # Cleanup
        output_path.unlink()
        doc_with_paragraphs.unlink()

    def test_heading_spacing(self, doc_with_paragraphs):
        """Headings should have 12pt before, 6pt after."""
        from pir_docx_normalizer import normalize_document
        from docx.shared import Pt

        output_path = doc_with_paragraphs.with_name("output.docx")
        normalize_document(doc_with_paragraphs, output_path, verbose=False)

        doc = Document(str(output_path))

        heading_paras = [p for p in doc.paragraphs
                        if p.style and 'Heading' in p.style.name and p.text.strip()]
        assert len(heading_paras) > 0, "Should have Heading paragraphs to test"

        for para in heading_paras:
            space_before = para.paragraph_format.space_before
            space_after = para.paragraph_format.space_after

            assert space_before is not None, f"Heading '{para.text[:20]}' space_before should be set"
            assert space_before.pt == self.HEADING_SPACE_BEFORE, \
                f"Heading space_before should be {self.HEADING_SPACE_BEFORE}pt, got {space_before.pt}pt"
            assert space_after is not None, f"Heading '{para.text[:20]}' space_after should be set"
            assert space_after.pt == self.HEADING_SPACE_AFTER, \
                f"Heading space_after should be {self.HEADING_SPACE_AFTER}pt, got {space_after.pt}pt"

        # Cleanup
        output_path.unlink()
        doc_with_paragraphs.unlink()

    def test_list_item_spacing(self, doc_with_paragraphs):
        """List items should have tight 2pt spacing."""
        from pir_docx_normalizer import normalize_document
        from docx.shared import Pt

        output_path = doc_with_paragraphs.with_name("output.docx")
        normalize_document(doc_with_paragraphs, output_path, verbose=False)

        doc = Document(str(output_path))

        list_paras = [p for p in doc.paragraphs
                     if p.style and 'List' in p.style.name and p.text.strip()]
        assert len(list_paras) > 0, "Should have List paragraphs to test"

        for para in list_paras:
            space_before = para.paragraph_format.space_before
            space_after = para.paragraph_format.space_after

            assert space_before is not None, "List space_before should be set"
            assert space_before.pt == self.LIST_SPACE, \
                f"List space_before should be {self.LIST_SPACE}pt, got {space_before.pt}pt"
            assert space_after is not None, "List space_after should be set"
            assert space_after.pt == self.LIST_SPACE, \
                f"List space_after should be {self.LIST_SPACE}pt, got {space_after.pt}pt"

        # Cleanup
        output_path.unlink()
        doc_with_paragraphs.unlink()

    def test_table_cell_spacing(self, doc_with_paragraphs):
        """Table cell paragraphs should have minimal 2pt spacing."""
        from pir_docx_normalizer import normalize_document
        from docx.shared import Pt

        output_path = doc_with_paragraphs.with_name("output.docx")
        normalize_document(doc_with_paragraphs, output_path, verbose=False)

        doc = Document(str(output_path))

        assert len(doc.tables) > 0, "Should have tables to test"

        cells_checked = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            space_before = para.paragraph_format.space_before
                            space_after = para.paragraph_format.space_after

                            assert space_before is not None, "Table cell space_before should be set"
                            assert space_before.pt == self.TABLE_CELL_SPACE, \
                                f"Table cell space_before should be {self.TABLE_CELL_SPACE}pt, got {space_before.pt}pt"
                            assert space_after is not None, "Table cell space_after should be set"
                            assert space_after.pt == self.TABLE_CELL_SPACE, \
                                f"Table cell space_after should be {self.TABLE_CELL_SPACE}pt, got {space_after.pt}pt"
                            cells_checked += 1

        assert cells_checked > 0, "Should have checked at least one table cell"

        # Cleanup
        output_path.unlink()
        doc_with_paragraphs.unlink()


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
